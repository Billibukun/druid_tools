import sqlite3
import pandas as pd
import logging
from datetime import datetime, timedelta
from pathlib import Path
from dataclasses import dataclass
from typing import Dict
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dateutil.relativedelta import relativedelta

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('birth_registration_report.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

@dataclass
class ReportConfig:
    """Configuration for report generation"""
    start_date: str
    end_date: str
    output_dir: str
    report_type: str  # "national", "state"
    period: str  # "monthly", "quarterly", "yearly"
    state: str = None
    include_charts: bool = True
    include_tables: bool = True
    save_data: bool = False

class ReportGenerator:
    """Comprehensive birth registration report generator"""
    def __init__(self, db_path: str):
        self.db_path = db_path
        self.conn = None
        self.queries = {
            "national_summary": """
            SELECT 
                COUNT(*) as total_registrations,
                COUNT(DISTINCT registration_center_state) as total_states,
                COUNT(DISTINCT registration_center_lga) as total_lgas,
                COUNT(DISTINCT registration_center) as total_centers,
                COUNT(CASE WHEN child_sex = 'Male' THEN 1 END) as male_children,
                COUNT(CASE WHEN child_sex = 'Female' THEN 1 END) as female_children,
                ROUND(AVG(mother_age_at_birth), 2) as avg_mother_age,
                ROUND(AVG(father_age_at_birth), 2) as avg_father_age,
                COUNT(CASE WHEN child_nin IS NOT NULL AND trim(child_nin) != '' THEN 1 END) as child_nin_count,
                COUNT(CASE WHEN mother_nin IS NOT NULL AND trim(mother_nin) != '' THEN 1 END) as mother_nin_count,
                COUNT(CASE WHEN father_nin IS NOT NULL AND trim(father_nin) != '' THEN 1 END) as father_nin_count
            FROM birth_records
            WHERE Date_Registerred BETWEEN :start_date AND :end_date
            """,

            "state_summary": """
            SELECT 
                registration_center_state,
                COUNT(*) as total_registrations,
                COUNT(DISTINCT registration_center_lga) as lga_count,
                COUNT(DISTINCT registration_center) as center_count,
                COUNT(CASE WHEN child_sex = 'Male' THEN 1 END) as male_count,
                COUNT(CASE WHEN child_sex = 'Female' THEN 1 END) as female_count,
                ROUND(AVG(JULIANDAY(Date_Registerred) - JULIANDAY(child_birth_date)), 2) as avg_registration_delay
            FROM birth_records
            WHERE Date_Registerred BETWEEN :start_date AND :end_date
            GROUP BY registration_center_state
            ORDER BY total_registrations DESC
            """,

            "state_center_stats": """
            SELECT 
                registration_center_state,
                COUNT(DISTINCT registration_center) as total_centers,
                ROUND(COUNT(*) * 1.0 / COUNT(DISTINCT registration_center), 2) as avg_registrations_per_center,
                COUNT(DISTINCT registration_center_lga) as total_lgas,
                ROUND(COUNT(*) * 1.0 / COUNT(DISTINCT Date_Registerred), 2) as avg_daily_registrations
            FROM birth_records
            WHERE Date_Registerred BETWEEN :start_date AND :end_date
            GROUP BY registration_center_state
            ORDER BY avg_registrations_per_center DESC
            """,

            "nin_coverage": """
            SELECT 
                registration_center_state,
                COUNT(*) as total_records,
                ROUND(COUNT(CASE WHEN child_nin IS NOT NULL AND trim(child_nin) != '' THEN 1 END) * 100.0 / COUNT(*), 2) as child_nin_pct,
                ROUND(COUNT(CASE WHEN mother_nin IS NOT NULL AND trim(mother_nin) != '' THEN 1 END) * 100.0 / COUNT(*), 2) as mother_nin_pct,
                ROUND(COUNT(CASE WHEN father_nin IS NOT NULL AND trim(father_nin) != '' THEN 1 END) * 100.0 / COUNT(*), 2) as father_nin_pct
            FROM birth_records
            WHERE Date_Registerred BETWEEN :start_date AND :end_date
            GROUP BY registration_center_state
            """,

            "data_quality": """
            SELECT 
                registration_center_state,
                COUNT(*) as total_records,
                ROUND(COUNT(CASE WHEN child_surname IS NULL OR trim(child_surname) = '' THEN 1 END) * 100.0 / COUNT(*), 2) as missing_surname_pct,
                ROUND(COUNT(CASE WHEN child_birth_date IS NULL THEN 1 END) * 100.0 / COUNT(*), 2) as missing_dob_pct,
                ROUND(COUNT(CASE WHEN mother_age_at_birth IS NULL THEN 1 END) * 100.0 / COUNT(*), 2) as missing_mother_age_pct,
                ROUND(COUNT(CASE WHEN father_age_at_birth IS NULL THEN 1 END) * 100.0 / COUNT(*), 2) as missing_father_age_pct
            FROM birth_records
            WHERE Date_Registerred BETWEEN :start_date AND :end_date
            GROUP BY registration_center_state
            """,

            "lga_summary": """
            SELECT 
                registration_center_lga,
                COUNT(*) as total_registrations,
                COUNT(DISTINCT registration_center) as center_count,
                COUNT(CASE WHEN child_sex = 'Male' THEN 1 END) as male_count,
                COUNT(CASE WHEN child_sex = 'Female' THEN 1 END) as female_count,
                COUNT(CASE WHEN child_sex = 'Male' AND JULIANDAY(:end_date) - JULIANDAY(child_birth_date) <= 365 THEN 1 END) as male_under_1,
                COUNT(CASE WHEN child_sex = 'Female' AND JULIANDAY(:end_date) - JULIANDAY(child_birth_date) <= 365 THEN 1 END) as female_under_1,
                COUNT(CASE WHEN child_sex = 'Male' AND JULIANDAY(:end_date) - JULIANDAY(child_birth_date) <= 1825 THEN 1 END) as male_under_5,
                COUNT(CASE WHEN child_sex = 'Female' AND JULIANDAY(:end_date) - JULIANDAY(child_birth_date) <= 1825 THEN 1 END) as female_under_5,
                ROUND(AVG(JULIANDAY(Date_Registerred) - JULIANDAY(child_birth_date)), 2) as avg_registration_delay,
                ROUND(AVG(mother_age_at_birth), 2) as avg_mother_age,
                ROUND(AVG(father_age_at_birth), 2) as avg_father_age
            FROM birth_records
            WHERE Date_Registerred BETWEEN :start_date AND :end_date AND registration_center_state = :state
            GROUP BY registration_center_lga
            ORDER BY total_registrations DESC
            """,

            "lga_center_stats": """
            SELECT 
                registration_center_lga,
                COUNT(DISTINCT registration_center) as total_centers,
                ROUND(COUNT(*) * 1.0 / COUNT(DISTINCT registration_center), 2) as avg_registrations_per_center,
                ROUND(COUNT(*) * 1.0 / COUNT(DISTINCT Date_Registerred), 2) as avg_daily_registrations
            FROM birth_records
            WHERE Date_Registerred BETWEEN :start_date AND :end_date AND registration_center_state = :state
            GROUP BY registration_center_lga
            ORDER BY avg_registrations_per_center DESC
            """,

            "lga_nin_coverage": """
            SELECT 
                registration_center_lga,
                COUNT(*) as total_records,
                ROUND(COUNT(CASE WHEN child_nin IS NOT NULL AND trim(child_nin) != '' THEN 1 END) * 100.0 / COUNT(*), 2) as child_nin_pct,
                ROUND(COUNT(CASE WHEN mother_nin IS NOT NULL AND trim(mother_nin) != '' THEN 1 END) * 100.0 / COUNT(*), 2) as mother_nin_pct,
                ROUND(COUNT(CASE WHEN father_nin IS NOT NULL AND trim(father_nin) != '' THEN 1 END) * 100.0 / COUNT(*), 2) as father_nin_pct
            FROM birth_records
            WHERE Date_Registerred BETWEEN :start_date AND :end_date AND registration_center_state = :state
            GROUP BY registration_center_lga
            """,

            "lga_data_quality": """
            SELECT 
                registration_center_lga,
                COUNT(*) as total_records,
                ROUND(COUNT(CASE WHEN child_surname IS NULL OR trim(child_surname) = '' THEN 1 END) * 100.0 / COUNT(*), 2) as missing_surname_pct,
                ROUND(COUNT(CASE WHEN child_birth_date IS NULL THEN 1 END) * 100.0 / COUNT(*), 2) as missing_dob_pct,
                ROUND(COUNT(CASE WHEN mother_age_at_birth IS NULL THEN 1 END) * 100.0 / COUNT(*), 2) as missing_mother_age_pct,
                ROUND(COUNT(CASE WHEN father_age_at_birth IS NULL THEN 1 END) * 100.0 / COUNT(*), 2) as missing_father_age_pct
            FROM birth_records
            WHERE Date_Registerred BETWEEN :start_date AND :end_date AND registration_center_state = :state
            GROUP BY registration_center_lga
            """
        }

    def connect(self):
        """Establish database connection"""
        try:
            self.conn = sqlite3.connect(self.db_path)
            logger.info(f"Connected to database: {self.db_path}")
        except sqlite3.Error as e:
            logger.error(f"Database connection error: {e}")
            raise

    def close(self):
        """Close database connection"""
        if self.conn:
            self.conn.close()
            logger.info("Database connection closed")

    def __enter__(self):
        self.connect()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()

    def _execute_query(self, query_name: str, params: Dict) -> pd.DataFrame:
        """Execute SQL query and return results"""
        try:
            query = self.queries.get(query_name)
            if not query:
                raise ValueError(f"Query '{query_name}' not found")
            return pd.read_sql_query(query, self.conn, params=params)
        except Exception as e:
            logger.error(f"Error executing query '{query_name}': {e}")
            raise

    def _add_report_header(self, doc: Document, config: ReportConfig):
        """Add report header and title page"""
        if config.report_type == "national":
            title = doc.add_heading('National Birth Registration Report', 0)
        else:
            title = doc.add_heading(f'{config.state} State Birth Registration Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        period = doc.add_paragraph()
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period.add_run(f"Reporting Period: {config.start_date} to {config.end_date}")

        timestamp = doc.add_paragraph()
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        doc.add_page_break()

    def _generate_national_summary(self, doc: Document, config: ReportConfig):
        """Generate national summary section"""
        doc.add_heading('National Summary', level=1)

        # National Summary
        query_name = 'national_summary'
        params = {
            'start_date': config.start_date,
            'end_date': config.end_date
        }
        df = self._execute_query(query_name, params)
        row = df.iloc[0]
        # Add summary paragraphs
        summary = doc.add_paragraph()
        summary.add_run("National Overview\n").bold = True
        summary.add_run(f"During the period {config.start_date} to {config.end_date}, ")
        summary.add_run(f"a total of {row['total_registrations']:,} birth registrations were recorded ")
        summary.add_run(f"across {row['total_states']} states and {row['total_lgas']:,} local government areas, ")
        summary.add_run(f"utilizing {row['total_centers']:,} registration centers.")

        # Gender distribution
        gender = doc.add_paragraph()
        gender.add_run("\nGender Distribution\n").bold = True
        total_registrations = row['total_registrations']
        male_pct = (row['male_children'] / total_registrations) * 100 if total_registrations else 0.0
        female_pct = (row['female_children'] / total_registrations) * 100 if total_registrations else 0.0
        gender.add_run(f"Male: {row['male_children']:,} ({male_pct:.1f}%)\n")
        gender.add_run(f"Female: {row['female_children']:,} ({female_pct:.1f}%)")

        # Age statistics
        age_stats = doc.add_paragraph()
        age_stats.add_run("\nParental Age Statistics\n").bold = True
        avg_mother_age = row['avg_mother_age'] if row['avg_mother_age'] is not None else 0.0 # Handle None
        avg_father_age = row['avg_father_age'] if row['avg_father_age'] is not None else 0.0 # Handle None
        age_stats.add_run(f"Average mother's age: {avg_mother_age:.1f} years\n")
        age_stats.add_run(f"Average father's age: {avg_father_age:.1f} years")

        # NIN coverage
        nin = doc.add_paragraph()
        nin.add_run("\nNational Identity Number (NIN) Coverage\n").bold = True
        total_registrations = row['total_registrations']
        child_nin_pct = (row['child_nin_count'] / total_registrations) * 100 if total_registrations else 0.0
        mother_nin_pct = (row['mother_nin_count'] / total_registrations) * 100 if total_registrations else 0.0
        father_nin_pct = (row['father_nin_count'] / total_registrations) * 100 if total_registrations else 0.0
        nin.add_run(f"Children: {row['child_nin_count']:,} ({child_nin_pct:.1f}%)\n")
        nin.add_run(f"Mothers: {row['mother_nin_count']:,} ({mother_nin_pct:.1f}%)\n")
        nin.add_run(f"Fathers: {row['father_nin_count']:,} ({father_nin_pct:.1f}%)")

        doc.add_paragraph()

    def _generate_state_summary(self, doc: Document, config: ReportConfig):
        """Generate state summary section"""

        if config.report_type == "national":
            doc.add_heading('State-Level Analysis', level=1)
            # Get state summary data
            df = self._execute_query('state_summary', {
                'start_date': config.start_date,
                'end_date': config.end_date
            })
            
            
            # Add  summary table
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Shading Accent 1'

            # Headers
            headers = table.rows[0].cells
            headers[0].text = 'State'
            headers[1].text = 'Registrations'
            headers[2].text = 'LGAs'
            headers[3].text = 'Centers'
            headers[4].text = 'Male/Female Ratio'
            headers[5].text = 'Avg Delay (days)'

            # Data rows
            for _, row in df.iterrows():
                cells = table.add_row().cells
                cells[0].text = row['registration_center_state']
                cells[1].text = f"{row['total_registrations']:,}"
                cells[2].text = str(row['lga_count'])
                cells[3].text = str(row['center_count'])
                gender_ratio = row['male_count'] / row['female_count'] if row['female_count'] > 0 else 0
                cells[4].text = f"{gender_ratio:.2f}"
                cells[5].text = f"{row['avg_registration_delay']:.1f}"

        
        else:
            doc.add_heading(f'LGA-Level Analysis - {config.state}', level=1)
            # Get LGA summary data
            df = self._execute_query('lga_summary', {
                'start_date': config.start_date,
                'end_date': config.end_date,
                'state': config.state
            })
            # Add  summary table
            table = doc.add_table(rows=1, cols=10)
            table.style = 'Light Shading Accent 1'

            # Headers
            headers = table.rows[0].cells
            headers[0].text = 'LGA'
            headers[1].text = 'Registrations'
            headers[2].text = 'Centers'
            headers[3].text = 'Male'
            headers[4].text = 'Female'
            headers[5].text = 'Male < 1'
            headers[6].text = 'Female < 1'
            headers[7].text = 'Male < 5'
            headers[8].text = 'Female < 5'
            headers[9].text = 'Avg Delay (days)'

            # Data rows
            for _, row in df.iterrows():
                cells = table.add_row().cells
                cells[0].text = row['registration_center_lga']
                cells[1].text = f"{row['total_registrations']:,}"
                cells[2].text = str(row['center_count'])
                cells[3].text = f"{row['male_count']:,}"
                cells[4].text = f"{row['female_count']:,}"
                cells[5].text = f"{row['male_under_1']:,}"
                cells[6].text = f"{row['female_under_1']:,}"
                cells[7].text = f"{row['male_under_5']:,}"
                cells[8].text = f"{row['female_under_5']:,}"
                cells[9].text = f"{row['avg_registration_delay']:.1f}"
            

        doc.add_paragraph()

    def _generate_nin_coverage(self, doc: Document, config: ReportConfig):
        """Generate NIN coverage analysis"""
        
        if config.report_type == 'national':
            doc.add_heading('NIN Coverage Analysis', level=1)
            df = self._execute_query('nin_coverage', {
                'start_date': config.start_date,
                'end_date': config.end_date
            })
        else:
            doc.add_heading(f'NIN Coverage Analysis - {config.state}', level=1)
            df = self._execute_query('lga_nin_coverage', {
                'start_date': config.start_date,
                'end_date': config.end_date,
                'state': config.state
            })
            

        # Add NIN coverage table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Shading Accent 1'

        # Headers
        headers = table.rows[0].cells
        if config.report_type == 'national':
            headers[0].text = 'State'
        else:
            headers[0].text = 'LGA'
        headers[1].text = 'Total Records'
        headers[2].text = 'Child NIN %'
        headers[3].text = 'Mother NIN %'
        headers[4].text = 'Father NIN %'

        # Data rows
        for _, row in df.iterrows():
            cells = table.add_row().cells
            if config.report_type == 'national':
                cells[0].text = row['registration_center_state']
            else:
                cells[0].text = row['registration_center_lga']
            cells[1].text = f"{row['total_records']:,}"
            cells[2].text = f"{row['child_nin_pct']:.1f}%"
            cells[3].text = f"{row['mother_nin_pct']:.1f}%"
            cells[4].text = f"{row['father_nin_pct']:.1f}%"

        doc.add_paragraph()

    def _generate_data_quality(self, doc: Document, config: ReportConfig):
        """Generate data quality analysis"""
        if config.report_type == 'national':
            doc.add_heading('Data Quality Analysis', level=1)
            df = self._execute_query('data_quality', {
                'start_date': config.start_date,
                'end_date': config.end_date
            })
        else:
            doc.add_heading(f'Data Quality Analysis - {config.state}', level=1)
            df = self._execute_query('lga_data_quality', {
                'start_date': config.start_date,
                'end_date': config.end_date,
                'state': config.state
            })
            

        # Calculate national averages
        national_stats = {
            'missing_surname': df['missing_surname_pct'].mean(),
            'missing_dob': df['missing_dob_pct'].mean(),
            'missing_mother_age': df['missing_mother_age_pct'].mean(),
            'missing_father_age': df['missing_father_age_pct'].mean()
        }

        # Add national summary
        summary = doc.add_paragraph()
        if config.report_type == "national":
             summary.add_run("National Data Quality Metrics\n").bold = True
        else:
             summary.add_run(f"{config.state} State Data Quality Metrics\n").bold = True
       
        summary.add_run(f"Missing Surname: {national_stats['missing_surname']:.2f}%\n")
        summary.add_run(f"Missing Date of Birth: {national_stats['missing_dob']:.2f}%\n")
        summary.add_run(f"Missing Mother's Age: {national_stats['missing_mother_age']:.2f}%\n")
        summary.add_run(f"Missing Father's Age: {national_stats['missing_father_age']:.2f}%\n")

        # Add state-wise table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Shading Accent 1'

        # Headers
        headers = table.rows[0].cells
        if config.report_type == 'national':
            headers[0].text = 'State'
        else:
            headers[0].text = 'LGA'
        headers[1].text = 'Missing Surname %'
        headers[2].text = 'Missing DOB %'
        headers[3].text = 'Missing Mother Age %'
        headers[4].text = 'Missing Father Age %'

        # Data rows
        for _, row in df.sort_values('missing_surname_pct', ascending=True).iterrows():
            cells = table.add_row().cells
            if config.report_type == 'national':
                cells[0].text = row['registration_center_state']
            else:
                cells[0].text = row['registration_center_lga']
            cells[1].text = f"{row['missing_surname_pct']:.2f}%"
            cells[2].text = f"{row['missing_dob_pct']:.2f}%"
            cells[3].text = f"{row['missing_mother_age_pct']:.2f}%"
            cells[4].text = f"{row['missing_father_age_pct']:.2f}%"

        doc.add_paragraph()

    def _generate_performance(self, doc: Document, config: ReportConfig):
        """Generate performance analysis"""
        if config.report_type == "national":
            doc.add_heading('State Performance Analysis', level=1)
             # Get state center statistics
            df = self._execute_query('state_center_stats', {
                'start_date': config.start_date,
                'end_date': config.end_date
            })
        else:
            doc.add_heading(f'LGA Performance Analysis - {config.state}', level=1)
            # Get LGA center statistics
            df = self._execute_query('lga_center_stats', {
                'start_date': config.start_date,
                'end_date': config.end_date,
                'state': config.state
            })
           

        # Add performance summary
        summary = doc.add_paragraph()
        if config.report_type == 'national':
            summary.add_run("Registration Center Performance\n").bold = True
        else:
            summary.add_run("LGA Registration Center Performance\n").bold = True

        national_avg = df['avg_registrations_per_center'].mean()

        if not df.empty:  # Check if DataFrame is not empty
            if config.report_type == 'national':
                top_entity = df.iloc[0]
                summary.add_run(f"National average registrations per center: {national_avg:.0f}\n")
                summary.add_run(f"Highest performing state: {top_entity['registration_center_state']} ")
                summary.add_run(f"({top_entity['avg_registrations_per_center']:.0f} registrations per center)\n")
            else:
                top_entity = df.iloc[0]
                summary.add_run(f"{config.state} average registrations per center: {national_avg:.0f}\n")
                summary.add_run(f"Highest performing LGA: {top_entity['registration_center_lga']} ")
                summary.add_run(f"({top_entity['avg_registrations_per_center']:.0f} registrations per center)\n")
        else:
            summary.add_run(f"National average registrations per center: {national_avg:.0f}\n")
            if config.report_type == 'national':
                summary.add_run("No state performance data available for this period.\n")
            else:
                summary.add_run(f"No LGA performance data available in {config.state} for this period.\n")

        # Add performance table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'

        # Headers
        headers = table.rows[0].cells
        if config.report_type == 'national':
            headers[0].text = 'State'
        else:
            headers[0].text = 'LGA'
        headers[1].text = 'Centers'
        headers[2].text = 'Avg. Reg/Center'
        headers[3].text = 'Daily Avg. Reg'

        # Data rows
        for _, row in df.sort_values('avg_registrations_per_center', ascending=False).iterrows():
            cells = table.add_row().cells
            if config.report_type == 'national':
                cells[0].text = row['registration_center_state']
            else:
                cells[0].text = row['registration_center_lga']
            
            cells[1].text = f"{row['total_centers']:,}"
            cells[2].text = f"{row['avg_registrations_per_center']:,.0f}"
            cells[3].text = f"{row['avg_daily_registrations']:,.1f}"

        doc.add_paragraph()

    def generate_report(self, config: ReportConfig) -> str:
        """Generate comprehensive report"""
        try:
            # Initialize document
            doc = Document()

            # Generate report sections
            self._add_report_header(doc, config)
            if config.report_type == 'national':
                 self._generate_national_summary(doc, config)
            self._generate_state_summary(doc, config)
            self._generate_nin_coverage(doc, config)
            self._generate_data_quality(doc, config)
            self._generate_performance(doc, config)

            # Save report
            Path(config.output_dir).mkdir(parents=True, exist_ok=True)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            if config.report_type == "national":
                if config.period == "monthly":
                    filename = f"{config.output_dir}/National_BR_Report_{config.start_date}_{timestamp}.docx"
                elif config.period == "quarterly":
                    filename = f"{config.output_dir}/National_BR_Report_Q{get_quarter(config.start_date)}_{timestamp}.docx"
                else: # yearly
                    filename = f"{config.output_dir}/National_BR_Report_{get_year(config.start_date)}_{timestamp}.docx"
            else:
                state_dir = Path(config.output_dir) / config.state
                if config.period == "monthly":
                    month_name = get_month_name(config.start_date)
                    month_dir = state_dir / "monthly" / month_name
                    month_dir.mkdir(parents=True, exist_ok=True)
                    filename = f"{month_dir}/{config.state}_BR_Report_{month_name}_{timestamp}.docx"
                elif config.period == "quarterly":
                    quarter_dir = state_dir / "quarterly"
                    quarter_dir.mkdir(parents=True, exist_ok=True)
                    filename = f"{quarter_dir}/{config.state}_BR_Report_Q{get_quarter(config.start_date)}_{timestamp}.docx"
                else:  # yearly
                    yearly_dir = state_dir / "yearly"
                    yearly_dir.mkdir(parents=True, exist_ok=True)
                    filename = f"{yearly_dir}/{config.state}_BR_Report_{get_year(config.start_date)}_{timestamp}.docx"
            
            doc.save(filename)

            logger.info(f"Report generated: {filename}")
            return filename

        except Exception as e:
            logger.error(f"Error generating report: {e}")
            raise

def get_month_name(date_str: str) -> str:
    """Get the month name from a date string (YYYY-MM-DD)."""
    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
    return date_obj.strftime('%B')

def get_quarter(date_str: str) -> int:
    """Get the quarter (1-4) from a date string (YYYY-MM-DD)."""
    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
    return (date_obj.month - 1) // 3 + 1

def get_year(date_str: str) -> int:
    """Get the year from a date string (YYYY-MM-DD)."""
    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
    return date_obj.year

def generate_national_report(db_path: str, start_date: str, end_date: str, output_dir: str, period: str) -> str:
    """Helper function to generate national report"""
    config = ReportConfig(
        start_date=start_date,
        end_date=end_date,
        output_dir=output_dir,
        report_type="national",
        period=period,
        include_charts=True
    )

    with ReportGenerator(db_path) as generator:
        report_path = generator.generate_report(config)

        print(f"Report generated: {report_path}")
        return report_path

def generate_state_report(db_path: str, state: str, start_date: str, end_date: str, output_dir: str, period: str) -> None:
    """Helper function to generate state-specific report and LGA level reports"""
    
    # Generate state-level report
    config = ReportConfig(
        start_date=start_date,
        end_date=end_date,
        output_dir=output_dir,
        report_type="state",
        period=period,
        state=state,
        include_charts=True
    )

    with ReportGenerator(db_path) as generator:
        generator.generate_report(config)

def generate_reports_for_period(db_path: str, start_date_str: str, end_date_str: str, output_dir: str, period: str) -> None:
    """Generate reports for a given period (monthly, quarterly, or yearly)"""

    start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
    end_date = datetime.strptime(end_date_str, '%Y-%m-%d')

    current_date = start_date
    while current_date <= end_date:
        if period == "monthly":
            # Calculate start and end dates for the month
            period_start_date = current_date.strftime('%Y-%m-%d')
            period_end_date = (current_date + relativedelta(months=1) - timedelta(days=1)).strftime('%Y-%m-%d')

            # Generate national report
            generate_national_report(db_path, period_start_date, period_end_date, f"{output_dir}/national/monthly", period)

            # Generate state reports
            with ReportGenerator(db_path) as generator:
                states_df = generator._execute_query('state_summary', {'start_date': period_start_date, 'end_date': period_end_date})
                for state in states_df['registration_center_state'].unique():
                    generate_state_report(db_path, state, period_start_date, period_end_date, f"{output_dir}/state", period)

            # Move to the next month
            current_date += relativedelta(months=1)

        elif period == "quarterly":
            # Calculate start and end dates for the quarter
            quarter = get_quarter(current_date.strftime('%Y-%m-%d'))
            year = current_date.year
            period_start_date = datetime(year, (quarter - 1) * 3 + 1, 1).strftime('%Y-%m-%d')
            period_end_date = (datetime(year, (quarter) * 3, 1) + relativedelta(months=1) - timedelta(days=1)).strftime('%Y-%m-%d')
            
            
            # Generate national report
            generate_national_report(db_path, period_start_date, period_end_date, f"{output_dir}/national/quarterly", period)

            # Generate state reports
            with ReportGenerator(db_path) as generator:
                states_df = generator._execute_query('state_summary', {'start_date': period_start_date, 'end_date': period_end_date})
                for state in states_df['registration_center_state'].unique():
                    generate_state_report(db_path, state, period_start_date, period_end_date, f"{output_dir}/state", period)

            # Move to the next quarter
            current_date += relativedelta(months=3)

        elif period == "yearly":
            # Calculate start and end dates for the year
            period_start_date = current_date.strftime('%Y-01-01')
            period_end_date = current_date.strftime('%Y-12-31')

            # Generate national report
            generate_national_report(db_path, period_start_date, period_end_date, f"{output_dir}/national/yearly", period)

            # Generate state reports
            with ReportGenerator(db_path) as generator:
                states_df = generator._execute_query('state_summary', {'start_date': period_start_date, 'end_date': period_end_date})
                for state in states_df['registration_center_state'].unique():
                    generate_state_report(db_path, state, period_start_date, period_end_date, f"{output_dir}/state", period)

            # Move to the next year
            current_date += relativedelta(years=1)
        
        else:
            raise ValueError("Invalid period specified. Choose from 'monthly', 'quarterly', or 'yearly'.")

if __name__ == "__main__":
    db_path = "data/database/birth_records.db"  # Replace with your database path
    output_dir = "reports"
    start_date_str = "2022-01-01"  # Example start date
    end_date_str = "2024-12-31"  # Example end date

    # Generate monthly reports for the specified date range:
    generate_reports_for_period(db_path, start_date_str, end_date_str, output_dir, "monthly")

    # Generate quarterly reports for the specified date range:
    generate_reports_for_period(db_path, start_date_str, end_date_str, output_dir, "quarterly")

    # Generate yearly reports for the specified date range:
    generate_reports_for_period(db_path, start_date_str, end_date_str, output_dir, "yearly")