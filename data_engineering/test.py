# birth_analysis.py

import os
import sys
import logging
import sqlite3
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import requests
import json
import base64
from pathlib import Path
from typing import Dict, List, Tuple, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ProjectManager:
    """Manages project structure and resources"""
    
    def __init__(self, root_dir: Optional[str] = None):
        self.root_dir = root_dir or os.getcwd()
        self.date_str = datetime.now().strftime('%Y%m%d')
        self._setup_structure()
        self._setup_logging()
    
    def _setup_structure(self):
        """Creates project directory structure"""
        # Main directories
        self.dirs = {
            'data': self._create_dir('data/databases'),
            'logs': self._create_dir('logs'),
            'resources': self._create_dir(f'resources/{self.date_str}'),
            'reports': self._create_dir(f'reports/{self.date_str}')
        }
        
        # Resource subdirectories
        self.resource_dirs = {
            'images': self._create_dir(f'resources/{self.date_str}/images'),
            'contexts': self._create_dir(f'resources/{self.date_str}/contexts'),
            'charts': self._create_dir(f'resources/{self.date_str}/charts')
        }
        
        # Report subdirectories
        self.report_dirs = {
            'text': self._create_dir(f'reports/{self.date_str}/text'),
            'docs': self._create_dir(f'reports/{self.date_str}/docs'),
            'quality': self._create_dir(f'reports/{self.date_str}/quality'),
            'analysis': self._create_dir(f'reports/{self.date_str}/analysis')
        }
    
    def _create_dir(self, path: str) -> str:
        """Creates directory and returns path"""
        full_path = os.path.join(self.root_dir, path)
        os.makedirs(full_path, exist_ok=True)
        return full_path
    
    def _setup_logging(self):
        """Configures logging"""
        log_file = os.path.join(self.dirs['logs'], f'analysis_{self.date_str}.log')
        
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file),
                logging.StreamHandler(sys.stdout)
            ]
        )
        self.logger = logging.getLogger('BirthAnalysis')
        self.logger.info('Project structure initialized')

class DataManager:
    """Handles data loading and preprocessing"""
    
    def __init__(self, db_path: str):
        self.db_path = db_path
        self.logger = logging.getLogger('BirthAnalysis.DataManager')
    
    def load_data(self) -> pd.DataFrame:
        """Loads data from SQLite database"""
        try:
            query = """
            SELECT *
            FROM birth_records
            """
            with sqlite3.connect(self.db_path) as conn:
                df = pd.read_sql_query(query, conn)
                self.logger.info(f"Loaded {len(df)} records from database")
                return df
        except Exception as e:
            self.logger.error(f"Error loading data: {e}")
            raise

class DataQualityChecker:
    """Handles data quality validation"""
    
    def __init__(self):
        self.logger = logging.getLogger('BirthAnalysis.QualityChecker')
        self._setup_validation_rules()
    
    def _setup_validation_rules(self):
        """Sets up validation rules"""
        self.rules = {
            'name_validation': {
                'min_length': 2,
                'max_length': 50,
                'special_chars': r'[^a-zA-Z\s\'-]'
            },
            'age_validation': {
                'min_mother_age': 12,
                'max_mother_age': 70,
                'min_father_age': 15,
                'max_father_age': 90
            },
            'required_fields': [
                'child_firstname', 'child_surname', 
                'mother_firstname', 'mother_surname',
                'father_firstname', 'father_surname'
            ]
        }
    
    def check_quality(self, df: pd.DataFrame) -> Tuple[pd.DataFrame, Dict[str, Any]]:
        """Performs comprehensive quality checks"""
        quality_issues = {}
        
        # Initialize quality flags
        df['has_quality_issues'] = False
        
        # Name quality checks
        name_issues = self._check_names(df)
        quality_issues['name_issues'] = name_issues
        df.loc[name_issues['records_with_issues'], 'has_quality_issues'] = True
        
        # Age validation
        age_issues = self._check_ages(df)
        quality_issues['age_issues'] = age_issues
        df.loc[age_issues['records_with_issues'], 'has_quality_issues'] = True
        
        # Completeness checks
        completeness_issues = self._check_completeness(df)
        quality_issues['completeness_issues'] = completeness_issues
        df.loc[completeness_issues['records_with_issues'], 'has_quality_issues'] = True
        
        # Family consistency checks
        family_issues = self._check_family_consistency(df)
        quality_issues['family_issues'] = family_issues
        df.loc[family_issues['records_with_issues'], 'has_quality_issues'] = True
        
        self.logger.info(f"Quality check completed. Found {df['has_quality_issues'].sum()} records with issues")
        
        return df, quality_issues
    
    def _check_names(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Checks name quality"""
        issues = {
            'short_names': [],
            'long_names': [],
            'special_chars': [],
            'records_with_issues': pd.Series(False, index=df.index)
        }
        
        name_columns = [
            'child_firstname', 'child_surname',
            'mother_firstname', 'mother_surname',
            'father_firstname', 'father_surname'
        ]
        
        for col in name_columns:
            if col in df.columns:
                # Check length
                short_names = df[col].str.len() < self.rules['name_validation']['min_length']
                long_names = df[col].str.len() > self.rules['name_validation']['max_length']
                special_chars = df[col].str.contains(self.rules['name_validation']['special_chars'], na=False)
                
                issues['short_names'].extend(df[short_names][col].tolist())
                issues['long_names'].extend(df[long_names][col].tolist())
                issues['special_chars'].extend(df[special_chars][col].tolist())
                
                issues['records_with_issues'] |= (short_names | long_names | special_chars)
        
        return issues
    
    def _check_ages(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Validates ages"""
        issues = {
            'invalid_mother_age': [],
            'invalid_father_age': [],
            'records_with_issues': pd.Series(False, index=df.index)
        }
        
        # Mother age validation
        if 'mother_age_at_birth' in df.columns:
            invalid_mother = (
                (df['mother_age_at_birth'] < self.rules['age_validation']['min_mother_age']) |
                (df['mother_age_at_birth'] > self.rules['age_validation']['max_mother_age'])
            )
            issues['invalid_mother_age'] = df[invalid_mother]['mother_age_at_birth'].tolist()
            issues['records_with_issues'] |= invalid_mother
        
        # Father age validation
        if 'father_age_at_birth' in df.columns:
            invalid_father = (
                (df['father_age_at_birth'] < self.rules['age_validation']['min_father_age']) |
                (df['father_age_at_birth'] > self.rules['age_validation']['max_father_age'])
            )
            issues['invalid_father_age'] = df[invalid_father]['father_age_at_birth'].tolist()
            issues['records_with_issues'] |= invalid_father
        
        return issues
    
    def _check_completeness(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Checks data completeness"""
        issues = {
            'missing_required_fields': {},
            'records_with_issues': pd.Series(False, index=df.index)
        }
        
        for field in self.rules['required_fields']:
            if field in df.columns:
                missing = df[field].isna() | (df[field] == '')
                issues['missing_required_fields'][field] = missing.sum()
                issues['records_with_issues'] |= missing
        
        return issues
    
    def _check_family_consistency(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Checks family data consistency"""
        issues = {
            'address_mismatch': [],
            'records_with_issues': pd.Series(False, index=df.index)
        }
        
        # Group by mother_id and check address consistency
        if 'mother_id' in df.columns and 'mother_address' in df.columns:
            address_groups = df.groupby('mother_id')['mother_address'].nunique()
            inconsistent_families = address_groups[address_groups > 1].index
            
            address_mismatch = df['mother_id'].isin(inconsistent_families)
            issues['address_mismatch'] = df[address_mismatch]['mother_id'].tolist()
            issues['records_with_issues'] |= address_mismatch
        
        return issues

class DataAnalyzer:
    """Performs comprehensive data analysis"""
    
    def __init__(self, project_manager: ProjectManager):
        self.project = project_manager
        self.logger = logging.getLogger('BirthAnalysis.Analyzer')
    
    def analyze_data(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Performs complete analysis on clean data"""
        results = {}
        
        # Demographic analysis
        self.logger.info("Performing demographic analysis...")
        results['demographics'] = self._analyze_demographics(df)
        
        # Geographic analysis
        self.logger.info("Performing geographic analysis...")
        results['geographic'] = self._analyze_geographic_distribution(df)
        
        # Temporal analysis
        self.logger.info("Performing temporal analysis...")
        results['temporal'] = self._analyze_temporal_patterns(df)
        
        # Family analysis
        self.logger.info("Performing family analysis...")
        results['family'] = self._analyze_family_patterns(df)
        
        return results
    
    def _analyze_demographics(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Analyzes demographic patterns"""
        demographics = {}
        
        # Gender distribution
        if 'child_sex' in df.columns:
            demographics['gender_dist'] = df['child_sex'].value_counts().to_dict()
        
        # Parent age distribution
        if 'mother_age_at_birth' in df.columns:
            demographics['mother_age_stats'] = {
                'mean': df['mother_age_at_birth'].mean(),
                'median': df['mother_age_at_birth'].median(),
                'std': df['mother_age_at_birth'].std(),
                'distribution': df['mother_age_at_birth'].value_counts().to_dict()
            }
        
        if 'father_age_at_birth' in df.columns:
            demographics['father_age_stats'] = {
                'mean': df['father_age_at_birth'].mean(),
                'median': df['father_age_at_birth'].median(),
                'std': df['father_age_at_birth'].std(),
                'distribution': df['father_age_at_birth'].value_counts().to_dict()
            }
        
        return demographics
    
    def _analyze_geographic_distribution(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Analyzes geographic distribution"""
        geographic = {}
        
        # State-level distribution
        if 'registration_center_state' in df.columns:
            geographic['state_dist'] = df['registration_center_state'].value_counts().to_dict()
        
        # LGA-level distribution
        if 'registration_center_lga' in df.columns:
            geographic['lga_dist'] = df['registration_center_lga'].value_counts().to_dict()
        
        return geographic
    
    def _analyze_temporal_patterns(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Analyzes temporal patterns"""
        temporal = {}
        
        # Registration patterns
        if 'Date_Registerred' in df.columns:
            df['Date_Registerred'] = pd.to_datetime(df['Date_Registerred'])
            temporal['monthly_registrations'] = (
                df.groupby(df['Date_Registerred'].dt.to_period('M'))
                .size()
                .to_dict()
            )
        
        return temporal
    
    def _analyze_family_patterns(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Analyzes family patterns"""
        family = {}
        
        # Family size distribution
        if 'mother_id' in df.columns:
            family_sizes = df.groupby('mother_id').size()
            family['size_distribution'] = family_sizes.value_counts().to_dict()
        
        # Multiple births
        if 'birth_type_desc' in df.columns:
            family['multiple_births'] = df['birth_type_desc'].value_counts().to_dict()
        
        return family

class ReportGenerator:
    """Generates comprehensive reports with AI insights"""
    
    def __init__(self, project_manager: ProjectManager):
        self.project = project_manager
        self.logger = logging.getLogger('BirthAnalysis.Reporter')
        self.ollama_url = "http://localhost:11434/api/generate"
    
    def generate_reports(
        self, 
        df: pd.DataFrame,
        quality_results: Dict[str, Any],
        analysis_results: Dict[str, Any]
        ):
        """Generates both quality and analysis reports"""
        # Generate quality report
        self.logger.info("Generating quality report...")
        self._generate_quality_report(df, quality_results)
        
        # Generate analysis report
        self.logger.info("Generating analysis report...")
        self._generate_analysis_report(df[~df['has_quality_issues']], analysis_results)
        
        # Generate executive summary
        self.logger.info("Generating executive summary...")
        self._generate_executive_summary(df, quality_results, analysis_results)
        
        self.logger.info("Report generation completed")

    def _generate_quality_report(self, df: pd.DataFrame, quality_results: Dict[str, Any]):
        """Generates detailed quality analysis report"""
        doc = Document()
        
        # Title and Introduction
        title = doc.add_heading('Birth Registration Data Quality Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        summary_data = {
            'total_records': len(df),
            'clean_records': len(df[~df['has_quality_issues']]),
            'issue_records': len(df[df['has_quality_issues']]),
            'quality_rate': f"{(len(df[~df['has_quality_issues']]) / len(df) * 100):.2f}%"
        }
        summary_prompt = f"""
        As a data quality expert, provide a concise executive summary of the birth registration data quality:
        
        Total Records: {summary_data['total_records']}
        Clean Records: {summary_data['clean_records']}
        Records with Issues: {summary_data['issue_records']}
        Data Quality Rate: {summary_data['quality_rate']}
        
        Focus on:
        1. Overall data quality assessment
        2. Significance of the quality metrics
        3. Impact on data usability
        """
        summary_response = self._get_llm_analysis(summary_prompt)
        doc.add_paragraph(summary_response)
        
        # 1. Name Quality Analysis
        self._add_name_quality_section(doc, quality_results['name_issues'])
        
        # 2. Age Validation Analysis
        self._add_age_validation_section(doc, quality_results['age_issues'])
        
        # 3. Completeness Analysis
        self._add_completeness_section(doc, quality_results['completeness_issues'])
        
        # 4. Family Consistency Analysis
        self._add_family_consistency_section(doc, quality_results['family_issues'])
        
        # 5. State-wise Quality Analysis
        self._add_state_quality_section(doc, df)
        
        # Save report
        doc.save(os.path.join(self.project.report_dirs['quality'], 'quality_report.docx'))
    
    def _add_name_quality_section(self, doc: Document, name_issues: Dict[str, Any]):
        """Adds detailed name quality analysis section"""
        doc.add_heading('Name Quality Analysis', 1)
        
        # Create and add detailed tables
        tables = []
        
        # 1. Short Names Analysis
        short_names_df = pd.DataFrame({
            'Name': name_issues['short_names'],
            'Length': [len(name) for name in name_issues['short_names']]
        }).head(20)  # Show top 20 examples
        tables.append(('Short Names Analysis', short_names_df))
        
        # 2. Long Names Analysis
        long_names_df = pd.DataFrame({
            'Name': name_issues['long_names'],
            'Length': [len(name) for name in name_issues['long_names']]
        }).head(20)
        tables.append(('Long Names Analysis', long_names_df))
        
        # 3. Special Character Analysis
        special_chars_df = pd.DataFrame({
            'Name': name_issues['special_chars'],
            'Special Characters': [re.findall(r'[^a-zA-Z\s\'-]', name) for name in name_issues['special_chars']]
        }).head(20)
        tables.append(('Special Characters in Names', special_chars_df))
        
        # Add each table with LLM analysis
        for title, df_table in tables:
            doc.add_heading(title, 2)
            
            # Add table to document
            table = doc.add_table(rows=1, cols=len(df_table.columns))
            table.style = 'Table Grid'
            
            # Add headers
            for i, column in enumerate(df_table.columns):
                table.rows[0].cells[i].text = column
            
            # Add data
            for _, row in df_table.iterrows():
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = str(value)
            
            # Get LLM analysis for this table
            table_prompt = f"""
            Analyze this {title.lower()} table of birth registration data:
            
            {df_table.to_string()}
            
            Provide:
            1. Key patterns or issues observed
            2. Potential impact on data quality
            3. Specific recommendations for improvement
            
            Focus on practical insights and actionable recommendations.
            """
            table_analysis = self._get_llm_analysis(table_prompt)
            doc.add_paragraph(table_analysis)
    
    def _add_age_validation_section(self, doc: Document, age_issues: Dict[str, Any]):
        """Adds detailed age validation analysis section"""
        doc.add_heading('Age Validation Analysis', 1)
        
        # Create visualization of age distributions
        fig = self._create_age_distribution_chart(age_issues)
        chart_path = os.path.join(self.project.resource_dirs['charts'], 'age_distribution.png')
        fig.write_image(chart_path)
        
        # Add chart to document
        doc.add_picture(chart_path, width=Inches(6))
        
        # Get VLM analysis of the chart
        chart_analysis = self._get_vlm_analysis(
            chart_path,
            "Analyze this age distribution chart focusing on potential age-related data quality issues"
        )
        doc.add_paragraph(chart_analysis)
        
        # Add detailed age issue tables
        tables = []
        
        # 1. Invalid Mother Ages
        if age_issues['invalid_mother_age']:
            mother_age_df = pd.DataFrame({
                'Age': age_issues['invalid_mother_age']
            }).value_counts().reset_index()
            mother_age_df.columns = ['Age', 'Count']
            tables.append(('Invalid Mother Ages', mother_age_df))
        
        # 2. Invalid Father Ages
        if age_issues['invalid_father_age']:
            father_age_df = pd.DataFrame({
                'Age': age_issues['invalid_father_age']
            }).value_counts().reset_index()
            father_age_df.columns = ['Age', 'Count']
            tables.append(('Invalid Father Ages', father_age_df))
        
        # Add each table with LLM analysis
        for title, df_table in tables:
            doc.add_heading(title, 2)
            
            # Add table
            table = doc.add_table(rows=1, cols=len(df_table.columns))
            table.style = 'Table Grid'
            
            # Add headers and data
            for i, column in enumerate(df_table.columns):
                table.rows[0].cells[i].text = column
            for _, row in df_table.iterrows():
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = str(value)
            
            # Get LLM analysis
            table_prompt = f"""
            Analyze this table of {title.lower()}:
            
            {df_table.to_string()}
            
            Provide:
            1. Assessment of the age patterns
            2. Potential data quality concerns
            3. Recommendations for validation improvements
            
            Focus on demographic feasibility and data integrity.
            """
            table_analysis = self._get_llm_analysis(table_prompt)
            doc.add_paragraph(table_analysis)
    
    def _add_completeness_section(self, doc: Document, completeness_issues: Dict[str, Any]):
        """Adds detailed completeness analysis section"""
        doc.add_heading('Data Completeness Analysis', 1)
        
        # Create completeness visualization
        missing_fields_df = pd.DataFrame.from_dict(
            completeness_issues['missing_required_fields'],
            orient='index',
            columns=['Missing Count']
        ).reset_index()
        missing_fields_df.columns = ['Field', 'Missing Count']
        
        fig = self._create_completeness_chart(missing_fields_df)
        chart_path = os.path.join(self.project.resource_dirs['charts'], 'completeness.png')
        fig.write_image(chart_path)
        
        # Add chart to document
        doc.add_picture(chart_path, width=Inches(6))
        
        # Get VLM analysis of the chart
        chart_analysis = self._get_vlm_analysis(
            chart_path,
            "Analyze this data completeness visualization focusing on patterns and critical gaps"
        )
        doc.add_paragraph(chart_analysis)
        
        # Add detailed completeness table
        doc.add_heading('Field-wise Completeness', 2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add headers
        headers = ['Field', 'Missing Records']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        
        # Add data
        for field, count in completeness_issues['missing_required_fields'].items():
            cells = table.add_row().cells
            cells[0].text = field
            cells[1].text = str(count)
        
        # Get LLM analysis
        table_prompt = f"""
        Analyze this data completeness table:
        
        {missing_fields_df.to_string()}
        
        Provide:
        1. Assessment of completeness patterns
        2. Critical fields with significant gaps
        3. Impact on data usability
        4. Recommendations for improving completeness
        
        Focus on practical implications and specific improvement strategies.
        """
        table_analysis = self._get_llm_analysis(table_prompt)
        doc.add_paragraph(table_analysis)
    
    def _add_family_consistency_section(self, doc: Document, family_issues: Dict[str, Any]):
        """Adds family consistency analysis section"""
        doc.add_heading('Family Data Consistency Analysis', 1)
        
        # Create family consistency visualization
        family_df = pd.DataFrame({
            'Family ID': family_issues['address_mismatch'],
            'Issue': 'Address Mismatch'
        })
        
        if not family_df.empty:
            fig = self._create_family_consistency_chart(family_df)
            chart_path = os.path.join(self.project.resource_dirs['charts'], 'family_consistency.png')
            fig.write_image(chart_path)
            
            # Add chart to document
            doc.add_picture(chart_path, width=Inches(6))
            
            # Get VLM analysis of the chart
            chart_analysis = self._get_vlm_analysis(
                chart_path,
                "Analyze this family consistency visualization focusing on patterns and anomalies"
            )
            doc.add_paragraph(chart_analysis)
            
            # Add detailed analysis
            inconsistent_counts = family_df.groupby('Issue').size().reset_index()
            inconsistent_counts.columns = ['Issue Type', 'Count']
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Add headers and data
            for i, column in enumerate(inconsistent_counts.columns):
                table.rows[0].cells[i].text = column
            for _, row in inconsistent_counts.iterrows():
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = str(value)
            
            # Get LLM analysis
            table_prompt = f"""
            Analyze this family consistency issues table:
            
            {inconsistent_counts.to_string()}
            
            Provide:
            1. Assessment of family data consistency
            2. Potential causes of inconsistencies
            3. Impact on family relationship analysis
            4. Recommendations for improving consistency
            
            Focus on data integrity and family record linkage.
            """
            table_analysis = self._get_llm_analysis(table_prompt)
            doc.add_paragraph(table_analysis)
    
    def _add_state_quality_section(self, doc: Document, df: pd.DataFrame):
        """Adds state-wise quality analysis section"""
        doc.add_heading('State-wise Quality Analysis', 1)
        
        # Calculate state-wise quality metrics
        state_metrics = df.groupby('registration_center_state').agg({
            'Birth_Reg_ID': 'count',
            'has_quality_issues': 'sum'
        }).reset_index()
        
        state_metrics['clean_rate'] = (
            (state_metrics['Birth_Reg_ID'] - state_metrics['has_quality_issues']) / 
            state_metrics['Birth_Reg_ID'] * 100
        ).round(2)
        
        state_metrics.columns = ['State', 'Total Records', 'Issues', 'Data Quality Rate (%)']
        
        # Create state-wise visualization
        fig = self._create_state_quality_chart(state_metrics)
        chart_path = os.path.join(self.project.resource_dirs['charts'], 'state_quality.png')
        fig.write_image(chart_path)
        
        # Add chart to document
        doc.add_picture(chart_path, width=Inches(6))
        
        # Get VLM analysis of the chart
        chart_analysis = self._get_vlm_analysis(
            chart_path,
            "Analyze this state-wise data quality visualization focusing on regional patterns and variations"
        )
        doc.add_paragraph(chart_analysis)
        
        # Add detailed state-wise table
        table = doc.add_table(rows=1, cols=len(state_metrics.columns))
        table.style = 'Table Grid'
        
        # Add headers and data
        for i, column in enumerate(state_metrics.columns):
            table.rows[0].cells[i].text = column
        for _, row in state_metrics.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        
        # Get LLM analysis
        table_prompt = f"""
        Analyze this state-wise data quality table:
        
        {state_metrics.to_string()}
        
        Provide:
        1. Assessment of regional quality patterns
        2. States with notable quality issues
        3. States with best practices
        4. Recommendations for improving state-level data quality
        
        Focus on regional variations and specific improvement strategies.
        """
        table_analysis = self._get_llm_analysis(table_prompt)
        doc.add_paragraph(table_analysis)

    def _get_llm_analysis(self, prompt: str) -> str:
        """Gets analysis from LLM"""
        try:
            response = requests.post(
                self.ollama_url,
                json={
                    "model": "nemotron-mini",
                    "prompt": prompt
                }
            )
            return response.json()['response']
        except Exception as e:
            self.logger.error(f"Error getting LLM analysis: {e}")
            return "Error generating analysis."

    def _get_vlm_analysis(self, image_path: str, prompt: str) -> str:
        """Gets analysis from VLM"""
        try:
            with open(image_path, 'rb') as img_file:
                image_base64 = base64.b64encode(img_file.read()).decode('utf-8')
            
            response = requests.post(
                self.ollama_url,
                json={
                    "model": "minicpm-v:latest",
                    "prompt": prompt,
                    "images": [image_base64]
                }
            )
            return response.json()['response']
        except Exception as e:
            self.logger.error(f"Error getting VLM analysis: {e}")
            return "Error analyzing visualization."

    def _create_age_distribution_chart(self, age_issues: Dict[str, Any]) -> go.Figure:
        """Creates age distribution visualization"""
        fig = make_subplots(rows=1, cols=2, subplot_titles=('Mother Age Issues', 'Father Age Issues'))
        
        # Mother age distribution
        if age_issues['invalid_mother_age']:
            mother_ages = pd.Series(age_issues['invalid_mother_age'])
            fig.add_trace(
                go.Histogram(x=mother_ages, name='Mother Ages', 
                           marker_color='#FF69B4', opacity=0.7),
                row=1, col=1
            )
        
        # Father age distribution
        if age_issues['invalid_father_age']:
            father_ages = pd.Series(age_issues['invalid_father_age'])
            fig.add_trace(
                go.Histogram(x=father_ages, name='Father Ages',
                           marker_color='#4169E1', opacity=0.7),
                row=1, col=2
            )
        
        fig.update_layout(
            title='Distribution of Invalid Ages',
            showlegend=True,
            template='plotly_white',
            height=500,
            width=1000
        )
        
        return fig

    def _create_completeness_chart(self, missing_fields_df: pd.DataFrame) -> go.Figure:
        """Creates completeness visualization"""
        fig = go.Figure()
        
        fig.add_trace(go.Bar(
            x=missing_fields_df['Field'],
            y=missing_fields_df['Missing Count'],
            marker_color='#2E8B57',
            opacity=0.7
        ))
        
        fig.update_layout(
            title='Missing Data by Field',
            xaxis_title='Field',
            yaxis_title='Number of Missing Records',
            template='plotly_white',
            height=500,
            width=1000,
            xaxis_tickangle=45
        )
        
        return fig

    def _create_family_consistency_chart(self, family_df: pd.DataFrame) -> go.Figure:
        """Creates family consistency visualization"""
        issue_counts = family_df['Issue'].value_counts()
        
        fig = go.Figure(data=[
            go.Pie(
                labels=issue_counts.index,
                values=issue_counts.values,
                hole=0.4,
                marker_colors=['#FF6B6B', '#4ECDC4', '#45B7D1']
            )
        ])
        
        fig.update_layout(
            title='Distribution of Family Data Inconsistencies',
            template='plotly_white',
            height=500,
            width=800
        )
        
        return fig

    def _create_state_quality_chart(self, state_metrics: pd.DataFrame) -> go.Figure:
        """Creates state-wise quality visualization"""
        fig = go.Figure()
        
        # Add bars for total records
        fig.add_trace(go.Bar(
            name='Total Records',
            x=state_metrics['State'],
            y=state_metrics['Total Records'],
            marker_color='#3498DB',
            opacity=0.7
        ))
        
        # Add bars for issues
        fig.add_trace(go.Bar(
            name='Records with Issues',
            x=state_metrics['State'],
            y=state_metrics['Issues'],
            marker_color='#E74C3C',
            opacity=0.7
        ))
        
        # Add line for quality rate
        fig.add_trace(go.Scatter(
            name='Data Quality Rate',
            x=state_metrics['State'],
            y=state_metrics['Data Quality Rate (%)'],
            yaxis='y2',
            line=dict(color='#2ECC71', width=3)
        ))
        
        fig.update_layout(
            title='State-wise Data Quality Overview',
            xaxis_title='State',
            yaxis_title='Number of Records',
            yaxis2=dict(
                title='Data Quality Rate (%)',
                overlaying='y',
                side='right',
                range=[0, 100]
            ),
            template='plotly_white',
            height=600,
            width=1200,
            xaxis_tickangle=45,
            barmode='group',
            legend=dict(
                yanchor="top",
                y=0.99,
                xanchor="right",
                x=0.99
            )
        )
        
        return fig

def analyze_birth_records(db_path: str, output_dir: Optional[str] = None):
    """Main function to run the complete birth records analysis"""
    try:
        # Initialize project structure
        project = ProjectManager(output_dir)
        logger = project.logger
        
        logger.info("Starting birth records analysis...")
        
        # Load data
        data_manager = DataManager(db_path)
        df = data_manager.load_data()
        logger.info(f"Loaded {len(df)} records")
        
        # Check data quality
        quality_checker = DataQualityChecker()
        df, quality_results = quality_checker.check_quality(df)
        logger.info("Completed quality checks")
        
        # Perform analysis on clean data
        clean_df = df[~df['has_quality_issues']].copy()
        analyzer = DataAnalyzer(project)
        analysis_results = analyzer.analyze_data(clean_df)
        logger.info("Completed data analysis")
        
        # Generate reports
        report_gen = ReportGenerator(project)
        report_gen.generate_reports(df, quality_results, analysis_results)
        logger.info("Generated reports")
        
        logger.info("Analysis completed successfully")
        
    except Exception as e:
        logger.error(f"Error in birth records analysis: {e}")
        raise

if __name__ == "__main__":
    # Example usage
    analyze_birth_records("birth_records.db")