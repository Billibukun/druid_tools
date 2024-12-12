# core/validation_checks.py

from typing import Dict
import logging

class ValidationChecks:
    """
    Core validation checks for birth registration data.
    Defines and manages all data quality rules and checks.
    """
    
    def __init__(self):
        self.logger = logging.getLogger('validation_checks')
        
        # Validation periods and thresholds
        self.OUTSIDE_HOURS_START_DATE = '2023-10-02'
        self.THRESHOLDS = {
            'mother_min_age': 16,
            'mother_max_age': 50,
            'father_min_age': 20,
            'father_max_age': 55,
            'mother_father_age_gap': 3,
            'high_activity_threshold': 50,
            'work_hours_start': 7,
            'work_hours_end': 19,
            'registration_delay_days': 365
        }
        
        # Define error categories and their SQL checks
        self.ERROR_CHECKS = {
            # Age-based checks
            'age_checks': {
                'mother_underage': f"""
                    CASE 
                        WHEN mother_age_at_birth < {self.THRESHOLDS['mother_min_age']}
                        THEN 1 
                        ELSE 0 
                    END
                """,
                'mother_overage': f"""
                    CASE 
                        WHEN mother_age_at_birth > {self.THRESHOLDS['mother_max_age']}
                        THEN 1 
                        ELSE 0 
                    END
                """,
                'father_underage': f"""
                    CASE 
                        WHEN father_age_at_birth < {self.THRESHOLDS['father_min_age']}
                        THEN 1 
                        ELSE 0 
                    END
                """,
                'father_overage': f"""
                    CASE 
                        WHEN father_age_at_birth > {self.THRESHOLDS['father_max_age']}
                        THEN 1 
                        ELSE 0 
                    END
                """,
                'age_gap': f"""
                    CASE 
                        WHEN (mother_age_at_birth - father_age_at_birth) > {self.THRESHOLDS['mother_father_age_gap']}
                        THEN 1 
                        ELSE 0 
                    END
                """
            },
            
            # Time-based checks
            'time_checks': {
                'outside_hours': f"""
                    CASE 
                        WHEN initiated_at >= '{self.OUTSIDE_HOURS_START_DATE}'
                        AND initiated_at IS NOT NULL 
                        AND (
                            CAST(strftime('%H', initiated_at) AS INTEGER) < {self.THRESHOLDS['work_hours_start']}
                            OR CAST(strftime('%H', initiated_at) AS INTEGER) > {self.THRESHOLDS['work_hours_end']}
                        )
                        THEN 1 
                        ELSE 0 
                    END
                """,
                'registration_delay': f"""
                    CASE 
                        WHEN JULIANDAY(Date_Registerred) - JULIANDAY(child_birth_date) > {self.THRESHOLDS['registration_delay_days']}
                        THEN 1 
                        ELSE 0 
                    END
                """
            },
            
            # Activity-based checks
            'activity_checks': {
                'high_daily_activity': f"""
                    Birth_Reg_ID IN (
                        SELECT Birth_Reg_ID 
                        FROM birth_records b2 
                        WHERE b2.registered_by_nin = birth_records.registered_by_nin
                        AND DATE(b2.initiated_at) = DATE(birth_records.initiated_at)
                        GROUP BY b2.registered_by_nin, DATE(b2.initiated_at)
                        HAVING COUNT(*) > {self.THRESHOLDS['high_activity_threshold']}
                    )
                """
            },
            
            # Completeness checks
            'completeness_checks': {
                'missing_child_name': """
                    CASE 
                        WHEN child_surname IS NULL 
                        OR trim(child_surname) = '' 
                        OR child_firstname IS NULL 
                        OR trim(child_firstname) = ''
                        THEN 1 
                        ELSE 0 
                    END
                """,
                'missing_child_birth': """
                    CASE 
                        WHEN child_birth_date IS NULL 
                        THEN 1 
                        ELSE 0 
                    END
                """,
                'missing_mother_details': """
                    CASE 
                        WHEN (
                            mother_surname IS NULL 
                            OR trim(mother_surname) = ''
                            OR mother_firstname IS NULL 
                            OR trim(mother_firstname) = ''
                            OR mother_age_at_birth IS NULL
                        )
                        THEN 1 
                        ELSE 0 
                    END
                """
            }
        }
    
    def get_error_query(self, error_type: str = 'all') -> str:
        """Generate SQL for error checks based on type"""
        if error_type == 'all':
            # Combine all checks
            all_checks = []
            for category in self.ERROR_CHECKS.values():
                all_checks.extend(category.values())
            return " OR ".join(f"({check})" for check in all_checks)
        elif error_type in self.ERROR_CHECKS:
            # Return specific category checks
            return " OR ".join(f"({check})" for check in self.ERROR_CHECKS[error_type].values())
        else:
            raise ValueError(f"Unknown error type: {error_type}")

    def get_error_summary_query(self) -> str:
        """Generate query for error summary by state"""
        # Build individual error columns
        error_columns = []
        for category, checks in self.ERROR_CHECKS.items():
            for name, check in checks.items():
                error_columns.append(f"SUM({check}) as {name}")
        
        # Basic query structure
        query = f"""
        SELECT 
            registration_center_state as state,
            COUNT(*) as total_records,
            {", ".join(error_columns)},
            SUM(CASE 
                WHEN approval_status_desc = 'Approved' 
                AND ({self.get_error_query()})
                THEN 1 ELSE 0 END
            ) as errors_approved,
            SUM(CASE 
                WHEN approval_status_desc = 'Approved' 
                AND NOT ({self.get_error_query()})
                THEN 1 ELSE 0 END
            ) as clean_approved
        FROM birth_records
        GROUP BY registration_center_state
        """
        return query

    def get_registration_quality_query(self) -> str:
        """Generate query for registration quality metrics"""
        return f"""
        SELECT 
            registration_center_state as state,
            COUNT(*) as total_records,
            ROUND(
                SUM(CASE WHEN ({self.get_error_query()}) THEN 1 ELSE 0 END) * 100.0 / COUNT(*), 
                2
            ) as error_rate,
            ROUND(
                SUM(CASE WHEN approval_status_desc = 'Approved' AND NOT ({self.get_error_query()})
                    THEN 1 ELSE 0 END) * 100.0 / 
                SUM(CASE WHEN approval_status_desc = 'Approved' THEN 1 ELSE 0 END),
                2
            ) as clean_approval_rate
        FROM birth_records
        GROUP BY registration_center_state
        """


# core/report_manager.py

from pathlib import Path
from datetime import datetime
import logging
from typing import Dict, List
import shutil

class ReportManager:
    """
    Manages report organization, storage, and retrieval with specific folder structure:
    
    reports/
    ├── logs/
    ├── national/
    │   └── {date}/
    │       └── BR_REG_NATIONAL_DATE_TIME.docx
    └── states/
        └── {date}/
            └── report_1/
                    BR_REG_STATE_DATE_TIME.docx
    """
    
    def __init__(self, base_dir: str = "reports"):
        self.base_dir = Path(base_dir)
        self.today = datetime.now().strftime('%Y%m%d')
        self._setup_directory_structure()
        self._setup_logging()
    
    def _setup_directory_structure(self):
        """Create required directory structure"""
        self.dirs = {
            'root': self.base_dir,
            'logs': self.base_dir / 'logs',
            'national': self.base_dir / 'national' / self.today,
            'states': self.base_dir / 'states' / self.today
        }
        
        for dir_path in self.dirs.values():
            dir_path.mkdir(parents=True, exist_ok=True)
            
        # Create error logs directory
        (self.dirs['logs'] / 'errors').mkdir(exist_ok=True)
    
    def _setup_logging(self):
        """Setup logging configuration"""
        log_file = self.dirs['logs'] / f'report_generation_{self.today}.log'
        error_file = self.dirs['logs'] / 'errors' / f'errors_{self.today}.log'
        
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file),
                logging.FileHandler(error_file),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger('report_manager')
    
    def _get_report_number(self, date_dir: Path) -> int:
        """Get next report number for state reports"""
        existing_reports = list(date_dir.glob('report_*'))
        if not existing_reports:
            return 1
        numbers = [int(d.name.split('_')[1]) for d in existing_reports]
        return max(numbers) + 1
    
    def save_national_report(self, report_content: bytes) -> str:
        """Save national report with specified naming convention"""
        try:
            timestamp = datetime.now()
            date_str = timestamp.strftime('%Y%m%d')
            time_str = timestamp.strftime('%H%M%S')
            
            filename = f"BR_REG_NATIONAL_{date_str}_{time_str}.docx"
            filepath = self.dirs['national'] / filename
            
            with open(filepath, 'wb') as f:
                f.write(report_content)
            
            self.logger.info(f"Saved national report: {filepath}")
            return str(filepath)
            
        except Exception as e:
            self.logger.error(f"Error saving national report: {e}")
            raise
    
    def save_state_report(self, state: str, report_content: bytes) -> str:
        """Save state report with specified structure and naming"""
        try:
            timestamp = datetime.now()
            date_str = timestamp.strftime('%Y%m%d')
            time_str = timestamp.strftime('%H%M%S')
            
            # Create report directory
            report_num = self._get_report_number(self.dirs['states'])
            report_dir = self.dirs['states'] / f"report_{report_num}"
            report_dir.mkdir(exist_ok=True)
            
            filename = f"BR_REG_{state}_{date_str}_{time_str}.docx"
            filepath = report_dir / filename
            
            with open(filepath, 'wb') as f:
                f.write(report_content)
            
            self.logger.info(f"Saved state report: {filepath}")
            return str(filepath)
            
        except Exception as e:
            self.logger.error(f"Error saving state report: {e}")
            raise
    
    def get_reports_by_date(self, date_str: str, report_type: str = 'national') -> List[Dict]:
        """Get all reports for a specific date"""
        try:
            if report_type == 'national':
                report_dir = self.base_dir / 'national' / date_str
            else:
                report_dir = self.base_dir / 'states' / date_str
            
            if not report_dir.exists():
                return []
            
            reports = []
            if report_type == 'national':
                for report_file in report_dir.glob('BR_REG_NATIONAL_*.docx'):
                    reports.append({
                        'path': str(report_file),
                        'timestamp': datetime.strptime(
                            report_file.stem.split('_')[-2:][0], 
                            '%Y%m%d'
                        ).strftime('%Y-%m-%d')
                    })
            else:
                for report_dir in report_dir.glob('report_*'):
                    for report_file in report_dir.glob('BR_REG_*.docx'):
                        state = report_file.stem.split('_')[2]
                        reports.append({
                            'path': str(report_file),
                            'state': state,
                            'report_number': int(report_dir.name.split('_')[1]),
                            'timestamp': datetime.strptime(
                                report_file.stem.split('_')[-2:][0], 
                                '%Y%m%d'
                            ).strftime('%Y-%m-%d')
                        })
            
            return sorted(reports, key=lambda x: x['timestamp'])
            
        except Exception as e:
            self.logger.error(f"Error retrieving reports for {date_str}: {e}")
            raise
    
    def cleanup_old_reports(self, days_to_keep: int = 30):
        """Archive or delete reports older than specified days"""
        try:
            from datetime import datetime, timedelta
            cutoff_date = datetime.now() - timedelta(days=days_to_keep)
            
            # Function to handle directory cleanup
            def cleanup_directory(directory: Path):
                if not directory.exists():
                    return
                
                for date_dir in directory.iterdir():
                    if not date_dir.is_dir():
                        continue
                    
                    try:
                        dir_date = datetime.strptime(date_dir.name, '%Y%m%d')
                        if dir_date < cutoff_date:
                            shutil.rmtree(date_dir)
                            self.logger.info(f"Removed old reports: {date_dir}")
                    except ValueError:
                        continue
            
            # Clean national and state directories
            cleanup_directory(self.base_dir / 'national')
            cleanup_directory(self.base_dir / 'states')
            
        except Exception as e:
            self.logger.error(f"Error during cleanup: {e}")
            raise


# core/reporters/base_reporter.py

import sqlite3
import pandas as pd
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional
import logging
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class BaseReporter:
    """
    Base class for CRVS reporting system.
    Provides common functionality for both state and national reports.
    """
    
    def __init__(self, db_path: str, target_path: Optional[str] = None):
        self.db_path = db_path
        self.conn = None
        self.validator = ValidationChecks()
        self.targets = pd.read_csv(target_path) if target_path else None
        self.logger = logging.getLogger(self.__class__.__name__)
        
        # Core metric queries
        self.METRIC_QUERIES = {
            'registration_stats': """
                SELECT 
                    COUNT(*) as total_registrations,
                    COUNT(DISTINCT registered_by_nin) as total_registrars,
                    COUNT(DISTINCT registration_center) as total_centers,
                    COUNT(DISTINCT registration_center_lga) as total_lgas,
                    SUM(CASE WHEN child_sex = 'MALE' THEN 1 ELSE 0 END) as male_count,
                    SUM(CASE WHEN child_sex = 'FEMALE' THEN 1 ELSE 0 END) as female_count,
                    SUM(CASE 
                        WHEN approval_status_desc = 'Approved' THEN 1 
                        ELSE 0 END
                    ) as approved_count,
                    SUM(CASE 
                        WHEN approval_status_desc = 'Queried' THEN 1 
                        ELSE 0 END
                    ) as queried_count,
                    SUM(CASE 
                        WHEN approval_status_desc = 'Pending' 
                        OR approval_status_desc IS NULL THEN 1 
                        ELSE 0 END
                    ) as pending_count,
                    ROUND(AVG(JULIANDAY(Date_Registerred) - JULIANDAY(child_birth_date))) as avg_delay
                FROM birth_records
                {where_clause}
            """,
            
            'monthly_trend': """
                SELECT 
                    strftime('%Y-%m', Date_Registerred) as month,
                    COUNT(*) as registrations,
                    COUNT(DISTINCT registered_by_nin) as registrars,
                    SUM(CASE WHEN child_sex = 'MALE' THEN 1 ELSE 0 END) as male_count,
                    SUM(CASE WHEN child_sex = 'FEMALE' THEN 1 ELSE 0 END) as female_count,
                    SUM(CASE 
                        WHEN approval_status_desc = 'Approved' THEN 1 
                        ELSE 0 END
                    ) as approved_count,
                    ROUND(AVG(JULIANDAY(Date_Registerred) - JULIANDAY(child_birth_date))) as avg_delay
                FROM birth_records
                {where_clause}
                GROUP BY month
                ORDER BY month
            """,
            
            'center_performance': """
                SELECT 
                    registration_center,
                    registration_center_state,
                    registration_center_lga,
                    COUNT(*) as total_registrations,
                    COUNT(DISTINCT registered_by_nin) as registrars,
                    SUM(CASE 
                        WHEN approval_status_desc = 'Approved' THEN 1 
                        ELSE 0 END
                    ) as approved_count,
                    ROUND(AVG(JULIANDAY(Date_Registerred) - JULIANDAY(child_birth_date))) as avg_delay,
                    SUM(CASE 
                        WHEN ({validation_checks}) THEN 1 
                        ELSE 0 END
                    ) as error_count
                FROM birth_records
                {where_clause}
                GROUP BY registration_center, registration_center_state, registration_center_lga
                ORDER BY total_registrations DESC
            """
        }
    
    def connect(self):
        """Establish database connection"""
        try:
            self.conn = sqlite3.connect(self.db_path)
            self.logger.info(f"Connected to database: {self.db_path}")
        except sqlite3.Error as e:
            self.logger.error(f"Database connection error: {e}")
            raise
    
    def close(self):
        """Close database connection"""
        if self.conn:
            self.conn.close()
            self.conn = None
            self.logger.info("Database connection closed")
    
    def get_metric(self, metric_name: str, where_clause: str = "") -> pd.DataFrame:
        """Get specified metric data"""
        try:
            if metric_name not in self.METRIC_QUERIES:
                raise ValueError(f"Unknown metric: {metric_name}")
            
            query = self.METRIC_QUERIES[metric_name].format(
                where_clause=where_clause,
                validation_checks=self.validator.get_error_query()
            )
            
            return pd.read_sql_query(query, self.conn)
        except Exception as e:
            self.logger.error(f"Error getting metric {metric_name}: {e}")
            raise
    
    def get_error_summary(self, where_clause: str = "") -> pd.DataFrame:
        """Get error summary based on validation checks"""
        try:
            query = self.validator.get_error_summary_query()
            if where_clause:
                query = query.replace(
                    "GROUP BY", 
                    f"WHERE {where_clause} GROUP BY"
                )
            return pd.read_sql_query(query, self.conn)
        except Exception as e:
            self.logger.error(f"Error getting error summary: {e}")
            raise

    def get_quality_metrics(self, where_clause: str = "") -> pd.DataFrame:
        """Get quality metrics"""
        try:
            query = self.validator.get_registration_quality_query()
            if where_clause:
                query = query.replace(
                    "GROUP BY",
                    f"WHERE {where_clause} GROUP BY"
                )
            return pd.read_sql_query(query, self.conn)
        except Exception as e:
            self.logger.error(f"Error getting quality metrics: {e}")
            raise

    def compare_with_target(self, actual: int, state: str) -> Dict:
        """Compare actual registrations with target"""
        if self.targets is None:
            return None
        
        try:
            target = self.targets[self.targets['state'] == state]['target'].iloc[0]
            achievement = (actual / target * 100) if target > 0 else 0
            
            return {
                'state': state,
                'target': target,
                'actual': actual,
                'achievement': achievement,
                'gap': target - actual if target > actual else 0
            }
        except Exception as e:
            self.logger.error(f"Error comparing with target for {state}: {e}")
            raise

    def _add_header(self, doc: Document, title: str):
        """Add standard report header"""
        main_title = doc.add_heading(title, 0)
        main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add spacing

    def _add_summary_table(self, doc: Document, data: Dict, title: str = "Summary"):
        """Add standard summary table to document"""
        doc.add_heading(title, level=1)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'
        
        for key, value in data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
            
        doc.add_paragraph()  # Add spacing

    def _format_number(self, value: float, percentage: bool = False) -> str:
        """Format numbers for report display"""
        if percentage:
            return f"{value:.1f}%" if value != 0 else "0%"
        return f"{value:,}" if value != 0 else "0"

    def __enter__(self):
        self.connect()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()


# core/reporters/state_reporter.py

from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class StateReporter(BaseReporter):
    """State-specific birth registration reporting"""
    
    def __init__(self, db_path: str, target_path: str, state: str):
        super().__init__(db_path, target_path)
        self.state = state
        self.where_clause = f"WHERE registration_center_state = '{state}'"
        
        # Additional state-specific queries
        self.state_queries = {
            'lga_summary': f"""
                SELECT 
                    registration_center_lga as lga,
                    COUNT(*) as total_registrations,
                    COUNT(DISTINCT registered_by_nin) as registrars,
                    COUNT(DISTINCT registration_center) as centers,
                    SUM(CASE WHEN child_sex = 'MALE' THEN 1 ELSE 0 END) as male_count,
                    SUM(CASE WHEN child_sex = 'FEMALE' THEN 1 ELSE 0 END) as female_count,
                    SUM(CASE WHEN approval_status_desc = 'Approved' THEN 1 ELSE 0 END) as approved,
                    ROUND(AVG(JULIANDAY(Date_Registerred) - JULIANDAY(child_birth_date))) as avg_delay
                FROM birth_records
                WHERE registration_center_state = '{state}'
                GROUP BY registration_center_lga
                ORDER BY total_registrations DESC
            """,
            
            'lga_errors': f"""
                SELECT 
                    registration_center_lga as lga,
                    COUNT(*) as total_records,
                    SUM(CASE WHEN ({self.validator.get_error_query()}) THEN 1 ELSE 0 END) as total_errors,
                    SUM(CASE WHEN approval_status_desc = 'Approved' AND ({self.validator.get_error_query()})
                        THEN 1 ELSE 0 END) as errors_approved,
                    SUM(CASE WHEN approval_status_desc = 'Approved' AND NOT ({self.validator.get_error_query()})
                        THEN 1 ELSE 0 END) as clean_approved,
                    COUNT(DISTINCT registration_center) as centers,
                    COUNT(DISTINCT registered_by_nin) as registrars
                FROM birth_records
                WHERE registration_center_state = '{state}'
                GROUP BY registration_center_lga
            """
        }
    
    def generate_report(self) -> bytes:
        """Generate comprehensive state report"""
        try:
            doc = Document()
            self._add_header(doc, f'Birth Registration Report - {self.state}')
            
            # Add sections
            self._add_executive_summary(doc)
            self._add_lga_analysis(doc)
            self._add_error_analysis(doc)
            self._add_registrar_performance(doc)
            self._add_target_analysis(doc)
            self._add_trend_analysis(doc)
            
            # Convert to bytes for saving
            from io import BytesIO
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return doc_bytes.read()
            
        except Exception as e:
            self.logger.error(f"Error generating report for {self.state}: {e}")
            raise
    
    def _add_executive_summary(self, doc: Document):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        
        # Get basic stats
        stats = self.get_metric('registration_stats', self.where_clause)
        row = stats.iloc[0]
        
        # Summary paragraph
        summary = doc.add_paragraph()
        summary.add_run(f"Registration Statistics for {self.state} State:\n").bold = True
        summary.add_run(f"Total Registrations: ").bold = True
        summary.add_run(f"{row['total_registrations']:,}\n")
        summary.add_run(f"Registration Centers: ").bold = True
        summary.add_run(f"{row['total_centers']:,}\n")
        summary.add_run(f"Total Registrars: ").bold = True
        summary.add_run(f"{row['total_registrars']:,}\n")
        
        # Add approval status
        approval = doc.add_paragraph()
        approval.add_run("Registration Status:\n").bold = True
        approval.add_run(f"- Approved: {row['approved_count']:,} "
                        f"({row['approved_count']/row['total_registrations']*100:.1f}%)\n")
        approval.add_run(f"- Queried: {row['queried_count']:,} "
                        f"({row['queried_count']/row['total_registrations']*100:.1f}%)\n")
        approval.add_run(f"- Pending: {row['pending_count']:,} "
                        f"({row['pending_count']/row['total_registrations']*100:.1f}%)\n")
    
    def _add_lga_analysis(self, doc: Document):
        """Add LGA analysis section"""
        doc.add_heading('LGA Analysis', level=1)
        
        # Get LGA data
        df_lga = pd.read_sql_query(self.state_queries['lga_summary'], self.conn)
        
        # Add summary text
        summary = doc.add_paragraph()
        summary.add_run(f"Analysis across {len(df_lga)} Local Government Areas:\n").bold = True
        
        # Add LGA table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Shading Accent 1'
        
        # Headers
        headers = table.rows[0].cells
        headers[0].text = 'LGA'
        headers[1].text = 'Registrations'
        headers[2].text = 'Centers'
        headers[3].text = 'Registrars'
        headers[4].text = 'Approval Rate'
        headers[5].text = 'Avg Delay'
        
        # Data rows
        for _, row in df_lga.iterrows():
            cells = table.add_row().cells
            cells[0].text = row['lga']
            cells[1].text = self._format_number(row['total_registrations'])
            cells[2].text = self._format_number(row['centers'])
            cells[3].text = self._format_number(row['registrars'])
            cells[4].text = self._format_number(row['approved']/row['total_registrations']*100, True)
            cells[5].text = f"{row['avg_delay']:.0f} days"
    
    def _add_error_analysis(self, doc: Document):
        """Add error analysis section"""
        doc.add_heading('Data Quality Analysis', level=1)
        
        # Get error data
        df_errors = pd.read_sql_query(self.state_queries['lga_errors'], self.conn)
        
        # Add summary
        total_errors = df_errors['total_errors'].sum()
        total_records = df_errors['total_records'].sum()
        error_rate = (total_errors / total_records * 100) if total_records > 0 else 0
        
        summary = doc.add_paragraph()
        summary.add_run("Overall Data Quality:\n").bold = True
        summary.add_run(f"- Error Rate: {error_rate:.1f}%\n")
        summary.add_run(f"- Clean Approved Records: "
                       f"{df_errors['clean_approved'].sum():,}\n")
        summary.add_run(f"- Records with Errors: {total_errors:,}\n")

        # Add LGA error table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Shading Accent 1'
        
        # Headers
        headers = table.rows[0].cells
        headers[0].text = 'LGA'
        headers[1].text = 'Total Records'
        headers[2].text = 'Error Rate'
        headers[3].text = 'Clean Approved'
        headers[4].text = 'Errors Approved'
        
        # Data rows
        for _, row in df_errors.iterrows():
            cells = table.add_row().cells
            cells[0].text = row['lga']
            cells[1].text = self._format_number(row['total_records'])
            cells[2].text = self._format_number(
                row['total_errors']/row['total_records']*100 
                if row['total_records'] > 0 else 0, 
                True
            )
            cells[3].text = self._format_number(row['clean_approved'])
            cells[4].text = self._format_number(row['errors_approved'])
    
    def _add_target_analysis(self, doc: Document):
        """Add target analysis section"""
        if self.targets is None:
            return
        
        doc.add_heading('Target Achievement Analysis', level=1)
        
        # Get clean approvals
        error_summary = self.get_error_summary(self.where_clause)
        clean_approved = error_summary['clean_approved'].iloc[0]
        
        # Compare with target
        target_info = self.compare_with_target(clean_approved, self.state)
        
        # Add summary
        summary = doc.add_paragraph()
        summary.add_run("Target Analysis:\n").bold = True
        summary.add_run(f"- Target: {target_info['target']:,}\n")
        summary.add_run(f"- Clean Approved: {clean_approved:,}\n")
        summary.add_run(f"- Achievement: {target_info['achievement']:.1f}%\n")
        if target_info['gap'] > 0:
            summary.add_run(f"- Gap to Target: {target_info['gap']:,}\n")
    
    def _add_trend_analysis(self, doc: Document):
        """Add trend analysis section"""
        doc.add_heading('Registration Trends', level=1)
        
        # Get monthly trends
        df_trend = self.get_metric('monthly_trend', self.where_clause)
        
        # Add trend table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Shading Accent 1'
        
        # Headers
        headers = table.rows[0].cells
        headers[0].text = 'Month'
        headers[1].text = 'Registrations'
        headers[2].text = 'Male'
        headers[3].text = 'Female'
        headers[4].text = 'Approval Rate'
        
        # Data rows (last 12 months)
        for _, row in df_trend.tail(12).iterrows():
            cells = table.add_row().cells
            cells[0].text = row['month']
            cells[1].text = self._format_number(row['registrations'])
            cells[2].text = self._format_number(row['male_count'])
            cells[3].text = self._format_number(row['female_count'])
            cells[4].text = self._format_number(
                row['approved_count']/row['registrations']*100 
                if row['registrations'] > 0 else 0, 
                True
            )


# core/reporters/national_reporter.py
from datetime import datetime, timedelta
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class NationalReporter(BaseReporter):
    """National-level birth registration reporting"""
    def __init__(self, db_path: str, target_path: str):
        super().__init__(db_path, target_path)
        
        # Additional national-level queries
        self.national_queries = {
            'state_summary': """
                SELECT 
                    registration_center_state as state,
                    COUNT(*) as total_registrations,
                    COUNT(DISTINCT registered_by_nin) as registrars,
                    COUNT(DISTINCT registration_center) as centers,
                    COUNT(DISTINCT registration_center_lga) as lgas,
                    SUM(CASE WHEN child_sex = 'MALE' THEN 1 ELSE 0 END) as male_count,
                    SUM(CASE WHEN child_sex = 'FEMALE' THEN 1 ELSE 0 END) as female_count,
                    SUM(CASE WHEN approval_status_desc = 'Approved' THEN 1 ELSE 0 END) as approved,
                    ROUND(AVG(JULIANDAY(Date_Registerred) - JULIANDAY(child_birth_date))) as avg_delay
                FROM birth_records
                GROUP BY registration_center_state
                ORDER BY total_registrations DESC
            """,
            
            'daily_stats': """
                SELECT 
                    DATE(Date_Registerred) as reg_date,
                    COUNT(*) as registrations,
                    COUNT(DISTINCT registration_center_state) as active_states,
                    COUNT(DISTINCT registered_by_nin) as active_registrars,
                    COUNT(DISTINCT registration_center) as active_centers,
                    SUM(CASE WHEN approval_status_desc = 'Approved' THEN 1 ELSE 0 END) as approved
                FROM birth_records
                WHERE Date_Registerred >= date('now', '-30 days')
                GROUP BY reg_date
                ORDER BY reg_date DESC
            """,
            
            'top_centers': """
                SELECT 
                    registration_center,
                    registration_center_state,
                    COUNT(*) as total_registrations,
                    COUNT(DISTINCT registered_by_nin) as registrars,
                    SUM(CASE WHEN approval_status_desc = 'Approved' THEN 1 ELSE 0 END) as approved,
                    ROUND(AVG(JULIANDAY(Date_Registerred) - JULIANDAY(child_birth_date))) as avg_delay,
                    SUM(CASE WHEN ({validation_checks}) THEN 1 ELSE 0 END) as error_count
                FROM birth_records
                GROUP BY registration_center, registration_center_state
                HAVING total_registrations > 1000
                ORDER BY total_registrations DESC
                LIMIT 20
            """
        }
    
    def get_all_states(self) -> List[str]:
        """Get list of all states"""
        query = "SELECT DISTINCT registration_center_state FROM birth_records ORDER BY registration_center_state"
        return pd.read_sql_query(query, self.conn)['registration_center_state'].tolist()
    
    def generate_report(self) -> bytes:
        """Generate comprehensive national report"""
        try:
            doc = Document()
            self._add_header(doc, 'National Birth Registration Report')
            
            # Add sections
            self._add_executive_summary(doc)
            self._add_state_performance(doc)
            self._add_error_analysis(doc)
            self._add_target_progress(doc)
            self._add_top_performers(doc)
            self._add_trend_analysis(doc)
            
            # Convert to bytes
            from io import BytesIO
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return doc_bytes.read()
            
        except Exception as e:
            self.logger.error(f"Error generating national report: {e}")
            raise
    
    def _add_executive_summary(self, doc: Document):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        
        # Get national stats
        stats = self.get_metric('registration_stats')
        row = stats.iloc[0]
        
        # Get state count
        states = self.get_all_states()
        
        # Summary paragraph
        summary = doc.add_paragraph()
        summary.add_run("National Overview\n").bold = True
        summary.add_run(f"A total of ")
        summary.add_run(f"{row['total_registrations']:,}").bold = True
        summary.add_run(" birth registrations were recorded across ")
        summary.add_run(f"{len(states)}").bold = True
        summary.add_run(" states, involving ")
        summary.add_run(f"{row['total_registrars']:,}").bold = True
        summary.add_run(" registrars operating in ")
        summary.add_run(f"{row['total_centers']:,}").bold = True
        summary.add_run(" registration centers.\n\n")
        
        # Registration status
        status = doc.add_paragraph()
        status.add_run("Registration Status:\n").bold = True
        status.add_run(f"- Approved: {row['approved_count']:,} "
                      f"({row['approved_count']/row['total_registrations']*100:.1f}%)\n")
        status.add_run(f"- Queried: {row['queried_count']:,} "
                      f"({row['queried_count']/row['total_registrations']*100:.1f}%)\n")
        status.add_run(f"- Pending: {row['pending_count']:,} "
                      f"({row['pending_count']/row['total_registrations']*100:.1f}%)\n")
        
        # Gender distribution
        gender = doc.add_paragraph()
        gender.add_run("Gender Distribution:\n").bold = True
        gender.add_run(f"- Male: {row['male_count']:,} "
                      f"({row['male_count']/row['total_registrations']*100:.1f}%)\n")
        gender.add_run(f"- Female: {row['female_count']:,} "
                      f"({row['female_count']/row['total_registrations']*100:.1f}%)\n")
    
    def _add_state_performance(self, doc: Document):
        """Add state performance analysis section"""
        doc.add_heading('State Performance Analysis', level=1)
        
        # Get state performance data
        df_states = pd.read_sql_query(self.national_queries['state_summary'], self.conn)
        
        # Summary text
        summary = doc.add_paragraph()
        top_state = df_states.iloc[0]
        summary.add_run("State Performance Overview:\n").bold = True
        summary.add_run(f"- Highest performing state: {top_state['state']} "
                       f"({top_state['total_registrations']:,} registrations)\n")
        
        # Add state comparison table
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Shading Accent 1'
        
        # Headers
        headers = table.rows[0].cells
        headers[0].text = 'State'
        headers[1].text = 'Registrations'
        headers[2].text = 'Centers'
        headers[3].text = 'LGAs'
        headers[4].text = 'Registrars'
        headers[5].text = 'Approval Rate'
        headers[6].text = 'Avg Delay'
        
        # Data rows
        for _, row in df_states.iterrows():
            cells = table.add_row().cells
            cells[0].text = row['state']
            cells[1].text = self._format_number(row['total_registrations'])
            cells[2].text = self._format_number(row['centers'])
            cells[3].text = self._format_number(row['lgas'])
            cells[4].text = self._format_number(row['registrars'])
            cells[5].text = self._format_number(
                row['approved']/row['total_registrations']*100 
                if row['total_registrations'] > 0 else 0, 
                True
            )
            cells[6].text = f"{row['avg_delay']:.0f} days"
    
    def _add_error_analysis(self, doc: Document):
        """Add national error analysis section"""
        doc.add_heading('Data Quality Analysis', level=1)
        
        # Get error summary
        error_summary = self.get_error_summary()
        
        # Calculate national totals
        total_records = error_summary['total_records'].sum()
        total_errors = error_summary['total_errors'].sum()
        error_rate = (total_errors / total_records * 100) if total_records > 0 else 0
        
        # Summary paragraph
        summary = doc.add_paragraph()
        summary.add_run("National Data Quality Overview:\n").bold = True
        summary.add_run(f"- Overall Error Rate: {error_rate:.1f}%\n")
        summary.add_run(f"- Total Records with Errors: {total_errors:,}\n")
        summary.add_run(f"- Clean Approved Records: "
                       f"{error_summary['clean_approved'].sum():,}\n")
        
        # Add state-wise quality table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Shading Accent 1'
        
        # Headers
        headers = table.rows[0].cells
        headers[0].text = 'State'
        headers[1].text = 'Total Records'
        headers[2].text = 'Error Rate'
        headers[3].text = 'Common Error'
        headers[4].text = 'Clean Records'
        headers[5].text = 'Error Approvals'
        
        # Data rows
        for _, row in error_summary.iterrows():
            cells = table.add_row().cells
            cells[0].text = row['state']
            cells[1].text = self._format_number(row['total_records'])
            
            error_rate = (row['total_errors'] / row['total_records'] * 100 
                         if row['total_records'] > 0 else 0)
            cells[2].text = self._format_number(error_rate, True)
            
            # Find most common error type
            error_counts = {
                'Age Issues': row['age_gap'] + row['mother_underage'] + 
                            row['mother_overage'] + row['father_underage'] + 
                            row['father_overage'],
                'Time Issues': row['outside_hours'],
                'Activity Issues': row['high_daily_activity'],
                'Missing Data': row['missing_child_name'] + row['missing_child_birth'] + 
                              row['missing_mother_details']
            }
            most_common = max(error_counts.items(), key=lambda x: x[1])
            cells[3].text = most_common[0]
            
            cells[4].text = self._format_number(row['clean_approved'])
            cells[5].text = self._format_number(row['errors_approved'])
    
    def _add_target_progress(self, doc: Document):
        """Add target progress section"""
        if self.targets is None:
            return
            
        doc.add_heading('Target Achievement Analysis', level=1)
        
        # Get clean approvals by state
        error_summary = self.get_error_summary()
        
        # Merge with targets
        progress_df = pd.merge(
            error_summary[['state', 'clean_approved']],
            self.targets,
            on='state'
        )
        
        # Calculate achievement
        progress_df['achievement'] = (
            progress_df['clean_approved'] / progress_df['target'] * 100).round(1)

        # Add national summary
        summary = doc.add_paragraph()
        total_target = progress_df['target'].sum()
        total_achieved = progress_df['clean_approved'].sum()
        national_achievement = (total_achieved / total_target * 100) if total_target > 0 else 0
        
        summary.add_run("National Target Progress:\n").bold = True
        summary.add_run(f"- National Target: {total_target:,}\n")
        summary.add_run(f"- Achieved: {total_achieved:,}\n")
        summary.add_run(f"- Achievement Rate: {national_achievement:.1f}%\n")
        
        # Add state progress table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        
        # Headers
        headers = table.rows[0].cells
        headers[0].text = 'State'
        headers[1].text = 'Target'
        headers[2].text = 'Achieved'
        headers[3].text = 'Achievement Rate'
        
        # Data rows
        for _, row in progress_df.sort_values('achievement', ascending=False).iterrows():
            cells = table.add_row().cells
            cells[0].text = row['state']
            cells[1].text = self._format_number(row['target'])
            cells[2].text = self._format_number(row['clean_approved'])
            cells[3].text = self._format_number(row['achievement'], True)
    
    def _add_top_performers(self, doc: Document):
        """Add top performing centres section."""
        doc.add_heading('Top Performing Centres', level=1)
    
        try:
            # Get centre performance data
            df_centres = pd.read_sql_query(
                self.national_queries['top_centers'].format(
                    validation_checks=self.validator.get_error_query()
                ),
                self.conn
            )
    
            if df_centres.empty:
                doc.add_paragraph("No data available for top performing centres.")
                return
    
            # Add introductory paragraph
            intro = doc.add_paragraph()
            intro.add_run(
                "The table below highlights the top-performing registration centres "
                "in terms of total registrations, with additional metrics such as "
                "error rate and average delay for each centre."
            ).italic = True
    
            # Add performance table
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Shading Accent 1'
    
            # Headers
            headers = table.rows[0].cells
            headers[0].text = 'Centre'
            headers[1].text = 'State'
            headers[2].text = 'Registrations'
            headers[3].text = 'Registrars'
            headers[4].text = 'Error Rate (%)'
            headers[5].text = 'Avg Delay (days)'
    
            # Data rows
            for _, row in df_centres.iterrows():
                cells = table.add_row().cells
                cells[0].text = row['registration_center']
                cells[1].text = row['registration_center_state']
                cells[2].text = self._format_number(row['total_registrations'])
                cells[3].text = self._format_number(row['registrars'])
                error_rate = (row['error_count'] / row['total_registrations'] * 100
                              if row['total_registrations'] > 0 else 0)
                cells[4].text = self._format_number(error_rate, True)
                cells[5].text = f"{row['avg_delay']:.0f}"
    
            # Add summary paragraph
            summary = doc.add_paragraph()
            summary.add_run(
                f"A total of {len(df_centres)} centres are listed as the top-performing centres. "
                "The data provides an overview of their performance, assisting in identifying "
                "centres with exemplary practices and potential for replication across other locations."
            )
    
        except Exception as e:
            # Add error message to the document in case of failure
            error_paragraph = doc.add_paragraph()
            error_paragraph.add_run("An error occurred while retrieving the top-performing centres:").bold = True
            error_paragraph.add_run(f" {e}")
            self.logger.error(f"Error in _add_top_performers: {e}")
    
    def _add_trend_analysis(self, doc: Document):
        """Add trend analysis section"""
        doc.add_heading('Registration Trends', level=1)
    
        # Get monthly trends
        df_trend = self.get_metric('monthly_trend')
    
        # Add trend table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Shading Accent 1'
    
        # Headers
        headers = table.rows[0].cells
        headers[0].text = 'Month'
        headers[1].text = 'Registrations'
        headers[2].text = 'Male'
        headers[3].text = 'Female'
        headers[4].text = 'Approval Rate'
    
        # Data rows (last 12 months)
        for _, row in df_trend.tail(12).iterrows():
            cells = table.add_row().cells
            cells[0].text = row['month']
            cells[1].text = self._format_number(row['registrations'])
            cells[2].text = self._format_number(row['male_count'])
            cells[3].text = self._format_number(row['female_count'])
            cells[4].text = self._format_number(
                row['approved_count']/row['registrations']*100
                if row['registrations'] > 0 else 0,
                True
            )


# core/report_scheduler.py
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.cron import CronTrigger
from datetime import datetime, time
import logging
from typing import Dict, List, Optional
from pathlib import Path


class ReportScheduler:
    """
    Handles automated scheduling of report generation
    """
    def __init__(self, reporting_system):
        self.system = reporting_system
        self.scheduler = BackgroundScheduler()
        self.logger = logging.getLogger('report_scheduler')
        
        # Setup logging
        self._setup_logging()
    
    def _setup_logging(self):
        """Setup scheduler-specific logging"""
        log_dir = Path("reports/logs/scheduler")
        log_dir.mkdir(parents=True, exist_ok=True)
        
        log_file = log_dir / f'scheduler_{datetime.now().strftime("%Y%m%d")}.log'
        
        handler = logging.FileHandler(log_file)
        handler.setFormatter(
            logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        )
        self.logger.addHandler(handler)
    
    def schedule_daily_reports(self, time_of_day: time = time(6, 0)):
        """Schedule comprehensive daily reports"""
        try:
            self.scheduler.add_job(
                self.system.generate_daily_reports,
                trigger=CronTrigger(
                    hour=time_of_day.hour, 
                    minute=time_of_day.minute
                ),
                id='daily_reports',
                name='Daily Reports Generation',
                misfire_grace_time=3600,  # 1 hour grace time
                coalesce=True,
                max_instances=1
            )
            
            self.logger.info(
                f"Scheduled daily reports for {time_of_day.strftime('%H:%M')}"
            )
        except Exception as e:
            self.logger.error(f"Error scheduling daily reports: {e}")
            raise
    
    def schedule_state_report(self, state: str, time_of_day: time):
        """Schedule report generation for specific state"""
        try:
            self.scheduler.add_job(
                self.system.generate_state_report,
                trigger=CronTrigger(
                    hour=time_of_day.hour, 
                    minute=time_of_day.minute
                ),
                args=[state],
                id=f'state_report_{state}',
                name=f'State Report - {state}',
                misfire_grace_time=1800,  # 30 minutes grace time
                coalesce=True
            )
            
            self.logger.info(
                f"Scheduled {state} report for {time_of_day.strftime('%H:%M')}"
            )
        except Exception as e:
            self.logger.error(f"Error scheduling state report for {state}: {e}")
            raise
    
    def schedule_national_report(self, time_of_day: time):
        """Schedule national report generation"""
        try:
            self.scheduler.add_job(
                self.system.generate_national_report,
                trigger=CronTrigger(
                    hour=time_of_day.hour, 
                    minute=time_of_day.minute
                ),
                id='national_report',
                name='National Report Generation',
                misfire_grace_time=3600,
                coalesce=True
            )
            
            self.logger.info(
                f"Scheduled national report for {time_of_day.strftime('%H:%M')}"
            )
        except Exception as e:
            self.logger.error(f"Error scheduling national report: {e}")
            raise
    
    def schedule_cleanup(self, days_to_keep: int = 30):
        """Schedule old report cleanup"""
        try:
            self.scheduler.add_job(
                self.system.cleanup_old_reports,
                trigger=CronTrigger(hour=1),  # Run at 1 AM
                args=[days_to_keep],
                id='cleanup',
                name='Report Cleanup',
                misfire_grace_time=3600,
                coalesce=True
            )
            
            self.logger.info(f"Scheduled cleanup to keep {days_to_keep} days of reports")
        except Exception as e:
            self.logger.error(f"Error scheduling cleanup: {e}")
            raise
    
    def get_schedule_status(self) -> List[Dict]:
        """Get status of all scheduled jobs"""
        jobs = []
        for job in self.scheduler.get_jobs():
            jobs.append({
                'id': job.id,
                'name': job.name,
                'next_run': job.next_run_time,
                'status': 'active' if job.next_run_time else 'paused'
            })
        return jobs
    
    def pause_job(self, job_id: str):
        """Pause specific scheduled job"""
        try:
            self.scheduler.pause_job(job_id)
            self.logger.info(f"Paused job: {job_id}")
        except Exception as e:
            self.logger.error(f"Error pausing job {job_id}: {e}")
            raise
    
    def resume_job(self, job_id: str):
        """Resume specific scheduled job"""
        try:
            self.scheduler.resume_job(job_id)
            self.logger.info(f"Resumed job: {job_id}")
        except Exception as e:
            self.logger.error(f"Error resuming job {job_id}: {e}")
            raise
    
    def start(self):
        """Start the scheduler"""
        try:
            self.scheduler.start()
            self.logger.info("Report scheduler started")
        except Exception as e:
            self.logger.error(f"Error starting scheduler: {e}")
            raise
    
    def stop(self):
        """Stop the scheduler"""
        try:
            self.scheduler.shutdown()
            self.logger.info("Report scheduler stopped")
        except Exception as e:
            self.logger.error(f"Error stopping scheduler: {e}")
            raise


# main
from typing import Dict, List, Optional
import logging
from pathlib import Path
from datetime import datetime, time

class CRVSReportingSystem:
    """
    Integrated CRVS Reporting System
    
    Combines:
    - Report generation (State and National)
    - File management and organization
    - Automated scheduling
    - Error tracking and logging
    """
    
    def __init__(self, 
                 db_path: str,
                 target_path: str,
                 base_dir: str = "reports"):
        self.db_path = db_path
        self.target_path = target_path
        self.report_manager = ReportManager(base_dir)
        self.scheduler = ReportScheduler(self)
        
        # Setup logging
        self._setup_logging()
    
    def _setup_logging(self):
        """Setup system-wide logging"""
        log_dir = Path("reports/logs/system")
        log_dir.mkdir(parents=True, exist_ok=True)
        
        log_file = log_dir / f'system_{datetime.now().strftime("%Y%m%d")}.log'
        
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger('crvs_system')
    
    def generate_state_report(self, state: str) -> str:
        """Generate and save state report"""
        try:
            self.logger.info(f"Starting report generation for {state}")
            
            # Generate report
            with StateReporter(self.db_path, self.target_path, state) as reporter:
                report_content = reporter.generate_report()
            
            # Save using report manager
            report_path = self.report_manager.save_state_report(state, report_content)
            
            self.logger.info(f"Completed report for {state}: {report_path}")
            return report_path
            
        except Exception as e:
            self.logger.error(f"Error generating report for {state}: {e}")
            raise
    
    def generate_national_report(self) -> str:
        """Generate and save national report"""
        try:
            self.logger.info("Starting national report generation")
            
            # Generate report
            with NationalReporter(self.db_path, self.target_path) as reporter:
                report_content = reporter.generate_report()
            
            # Save using report manager
            report_path = self.report_manager.save_national_report(report_content)
            
            self.logger.info(f"Completed national report: {report_path}")
            return report_path
            
        except Exception as e:
            self.logger.error(f"Error generating national report: {e}")
            raise
    
    def generate_daily_reports(self) -> Dict:
        """Generate complete set of daily reports"""
        try:
            self.logger.info("Starting daily report generation")
            
            results = {
                'timestamp': datetime.now().isoformat(),
                'national': None,
                'states': {}
            }
            
            # Generate national report
            try:
                national_path = self.generate_national_report()
                results['national'] = {
                    'status': 'success',
                    'path': national_path
                }
            except Exception as e:
                results['national'] = {
                    'status': 'failed',
                    'error': str(e)
                }
            
            # Generate state reports
            with NationalReporter(self.db_path, self.target_path) as reporter:
                states = reporter.get_all_states()
                
                for state in states:
                    try:
                        report_path = self.generate_state_report(state)
                        results['states'][state] = {
                            'status': 'success',
                            'path': report_path
                        }
                    except Exception as e:
                        results['states'][state] = {
                            'status': 'failed',
                            'error': str(e)
                        }
            
            self.logger.info("Completed daily report generation")
            return results
            
        except Exception as e:
            self.logger.error(f"Error in daily report generation: {e}")
            raise
    
    def setup_daily_schedule(self, 
                           national_time: time = time(7, 0),
                           state_time: time = time(6, 0),
                           cleanup_days: int = 30):
        """Setup complete daily schedule"""
        try:
            # Schedule national report
            self.scheduler.schedule_national_report(national_time)
            
            # Schedule state reports
            with NationalReporter(self.db_path, self.target_path) as reporter:
                states = reporter.get_all_states()
                for idx, state in enumerate(states):
                    # Stagger state reports by 5 minutes
                    report_time = time(
                        state_time.hour,
                        state_time.minute + (idx * 5) % 60,
                        state_time.second
                    )
                    self.scheduler.schedule_state_report(state, report_time)
            
            # Schedule cleanup
            self.scheduler.schedule_cleanup(cleanup_days)
            
            # Start scheduler
            self.scheduler.start()
            
        except Exception as e:
            self.logger.error(f"Error setting up schedule: {e}")
            raise
    
    def cleanup_old_reports(self, days_to_keep: int = 30):
        """Clean up old reports"""
        try:
            self.report_manager.cleanup_old_reports(days_to_keep)
        except Exception as e:
            self.logger.error(f"Error cleaning up reports: {e}")
            raise
    
    def get_schedule_status(self) -> List[Dict]:
        """Get current schedule status"""
        return self.scheduler.get_schedule_status()
    
    def stop_scheduler(self):
        """Stop the scheduler"""
        self.scheduler.stop()